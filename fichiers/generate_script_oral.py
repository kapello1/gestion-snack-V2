# -*- coding: utf-8 -*-
"""
Génération du Script Oral — Présentation TFE Gestion Snack
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Marges ───────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Couleurs ──────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1A, 0x23, 0x5E)
TEAL   = RGBColor(0x00, 0x7A, 0x8A)
ORANGE = RGBColor(0xE6, 0x7E, 0x22)
GREEN  = RGBColor(0x1E, 0x8B, 0x4C)
RED    = RGBColor(0xC0, 0x39, 0x2B)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
DGRAY  = RGBColor(0x33, 0x33, 0x33)
MGRAY  = RGBColor(0x77, 0x77, 0x77)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_shading(paragraph, hex_color):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    pPr.append(shd)

def page_break():
    doc.add_page_break()

def spacer(n=1):
    for _ in range(n):
        p = doc.add_paragraph('')
        p.paragraph_format.space_after = Pt(2)

def slide_header(number, title, duration, color=NAVY):
    """Bannière colorée pour chaque slide."""
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    # Numéro
    c0 = tbl.rows[0].cells[0]
    c0.text = f'SLIDE {number}'
    set_cell_bg(c0, '1A235E')
    for p in c0.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = WHITE
            r.font.size = Pt(11)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Titre
    c1 = tbl.rows[0].cells[1]
    c1.text = title
    set_cell_bg(c1, '007A8A')
    for p in c1.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = WHITE
            r.font.size = Pt(12)
    # Durée
    c2 = tbl.rows[0].cells[2]
    c2.text = f'⏱  {duration}'
    set_cell_bg(c2, 'E67E22')
    for p in c2.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = WHITE
            r.font.size = Pt(11)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Largeurs
    tbl.columns[0].width = Inches(1.2)
    tbl.columns[1].width = Inches(4.5)
    tbl.columns[2].width = Inches(1.5)
    doc.add_paragraph()

def spoken(text):
    """Texte à dire — corps principal."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.2)
    run = p.add_run(text)
    run.font.size = Pt(11.5)
    run.font.color.rgb = DGRAY
    return p

def spoken_bold(text):
    """Texte à dire avec emphase."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.2)
    run = p.add_run(text)
    run.font.size = Pt(11.5)
    run.font.bold = True
    run.font.color.rgb = NAVY
    return p

def action_note(text):
    """Encadré action de démo (fond vert clair)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_after  = Pt(4)
    add_shading(p, 'E8F8EE')
    run = p.add_run(f'▶  {text}')
    run.font.size   = Pt(10.5)
    run.font.italic = True
    run.font.color.rgb = GREEN
    return p

def tip_note(text):
    """Conseil / tip (fond orange clair)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_after  = Pt(4)
    add_shading(p, 'FEF3E2')
    run = p.add_run(f'💡  {text}')
    run.font.size   = Pt(10.5)
    run.font.italic = True
    run.font.color.rgb = ORANGE
    return p

def warning_note(text):
    """Mise en garde (fond rouge clair)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_after  = Pt(4)
    add_shading(p, 'FDEDEC')
    run = p.add_run(f'⚠  {text}')
    run.font.size   = Pt(10.5)
    run.font.italic = True
    run.font.color.rgb = RED
    return p

def section_title(text, color=NAVY):
    """Titre de section Q&R."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.size  = Pt(14)
    run.font.bold  = True
    run.font.color.rgb = color
    return p

def qa_question(text):
    """Question du jury."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.2)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    add_shading(p, 'EEF0F7')
    run = p.add_run(f'❓  {text}')
    run.font.size  = Pt(11.5)
    run.font.bold  = True
    run.font.color.rgb = NAVY
    return p

def qa_answer(text):
    """Réponse préparée."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.5)
    p.paragraph_format.space_after  = Pt(4)
    add_shading(p, 'F0F9F0')
    run = p.add_run(f'✔  {text}')
    run.font.size  = Pt(11)
    run.font.color.rgb = GREEN
    return p

# ══════════════════════════════════════════════════════════════════════════════
#  PAGE DE GARDE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SCRIPT ORAL')
run.font.size  = Pt(32)
run.font.bold  = True
run.font.color.rgb = NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Présentation TFE — Gestion Snack')
run.font.size  = Pt(18)
run.font.color.rgb = TEAL

doc.add_paragraph()

info = doc.add_table(rows=4, cols=2)
info.style = 'Table Grid'
pairs = [
    ('Étudiant',      'Tiegni Gamo Bernard-Joël'),
    ('Établissement', 'Institut Libre Marie Haps — Section Informatique'),
    ('Durée totale',  '≈ 20 minutes'),
    ('Date',          '17 juin 2026'),
]
for i, (k, v) in enumerate(pairs):
    info.rows[i].cells[0].text = k
    info.rows[i].cells[1].text = v
    set_cell_bg(info.rows[i].cells[0], '1A235E')
    for p2 in info.rows[i].cells[0].paragraphs:
        for r2 in p2.runs:
            r2.font.color.rgb = WHITE
            r2.font.bold = True
            r2.font.size = Pt(10.5)
    for p2 in info.rows[i].cells[1].paragraphs:
        for r2 in p2.runs:
            r2.font.size = Pt(10.5)

doc.add_paragraph()
tip_note("Ce document est votre script de présentation. Lisez-le à voix haute plusieurs fois "
         "pour mémoriser le rythme. Les encadrés verts ▶ indiquent les actions à faire "
         "sur l'écran pendant la présentation.")
warning_note("Ne lisez PAS ce script mot à mot face au jury. Appropriez-vous les idées et "
             "parlez naturellement. Ce texte est une base, pas une récitation.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SLIDE 1 — PAGE DE TITRE
# ══════════════════════════════════════════════════════════════════════════════
slide_header(1, 'Page de titre', '~40 secondes')

spoken("Bonjour à tous.")
spacer()
spoken("Je m'appelle Tiegni Gamo Bernard-Joël, et je vais vous présenter mon projet de TFE intitulé "
       "Gestion Snack — une application web complète pour la gestion d'un établissement de restauration rapide.")
spacer()
spoken("Cette application a été développée avec React pour le frontend, Spring Boot pour le backend, "
       "et PostgreSQL comme base de données. Elle est déployée en production dans le cloud.")
spacer()
spoken("La particularité de ce projet, c'est qu'il intègre une fonctionnalité que peu d'applications "
       "de ce type proposent : un chatbot vocal intelligent, qui combine la reconnaissance vocale, "
       "l'intelligence artificielle et la synthèse vocale naturelle.")
spacer()
spoken("Je vais vous présenter tout cela en environ vingt minutes.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SLIDE 2 — CONTEXTE
# ══════════════════════════════════════════════════════════════════════════════
slide_header(2, 'Contexte & Problématique', '~1 min 45')

spoken("Avant de parler de ma solution, parlons du problème.")
spacer()
spoken("Imaginons un snack qui fonctionne sans outil informatique. Le serveur prend les commandes "
       "sur un carnet, les crie en cuisine, le cuisinier note sur un tableau blanc. En caisse, "
       "on additionne à la main. Le stock ? On regarde dans les placards quand on pense à y penser. "
       "Et en fin de mois, impossible de savoir quels plats se sont le mieux vendus.")
spacer()
spoken("Ce scénario, c'est la réalité de beaucoup d'établissements. Et ces problèmes ont des "
       "conséquences directes : des commandes perdues, des ruptures de stock imprévues, des erreurs "
       "de caisse, et une impossibilité totale de prendre des décisions basées sur des données.")
spacer()
spoken("C'est exactement ce problème que mon application résout.")
spacer()
spoken("À droite de votre écran, vous voyez les six grandes fonctionnalités que j'ai développées "
       "pour y répondre : les commandes en temps réel par WebSocket, le décrément automatique du "
       "stock par des triggers PostgreSQL, une caisse multi-méthodes incluant le paiement en ligne "
       "Stripe, un tableau de bord analytique, les réservations en ligne, et le chatbot vocal.")
spacer()
spoken("Je vais maintenant vous montrer comment j'ai construit tout ça.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SLIDE 3 — ÉTUDE DE L'EXISTANT
# ══════════════════════════════════════════════════════════════════════════════
slide_header(3, "Étude de l'Existant", '~2 minutes')

spoken("Avant de concevoir quoi que ce soit, j'ai analysé ce qui existait déjà sur le marché.")
spacer()
spoken("J'ai étudié quatre solutions. Les trois premières sont Square POS, Lightspeed Restaurant, "
       "et Tiller Systems.")
spacer()
spoken("Ces solutions ont des qualités indéniables. Square POS est intuitif et gère bien les "
       "paiements. Lightspeed est très complet pour les restaurants multi-sites. Tiller a une belle "
       "interface tactile avec des rapports.")
spacer()
spoken("Mais j'ai identifié des lacunes communes à toutes ces solutions.")
spacer()
spoken("Premièrement, aucune n'intègre un chatbot vocal basé sur l'intelligence artificielle. "
       "Deuxièmement, leur gestion du stock est soit inexistante, soit basique — sans décrément "
       "automatique garanti par la base de données. Troisièmement, ces solutions sont soit trop "
       "chères, soit trop complexes pour une petite structure. Et enfin, elles ne proposent pas "
       "d'API REST documentée, ce qui empêche toute intégration personnalisée.")
spacer()
spoken("Mon projet se positionne sur ces quatre axes de différenciation. Et en particulier sur "
       "le chatbot vocal, qui est vraiment le point unique sur ce marché.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SLIDE 4 — ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
slide_header(4, 'Architecture Technique', '~2 minutes')

spoken("Parlons maintenant de comment l'application est construite.")
spacer()
spoken("J'ai choisi une architecture en trois couches, totalement déployée dans le cloud.")
spacer()
spoken("La première couche, c'est la présentation : le frontend React 19, buildé avec Vite 7 "
       "et stylisé avec Tailwind CSS. Il est hébergé sur Vercel, qui offre un CDN mondial — "
       "concrètement, l'application se charge rapidement depuis n'importe quel pays.")
spacer()
spoken("La deuxième couche, c'est la logique métier : le backend Spring Boot 3, écrit en Java 17. "
       "Il expose une API REST complète, gère les WebSockets pour le temps réel, et orchestre "
       "toutes les règles métier. Il tourne sur Render, dans un conteneur Linux — ce qui répond "
       "à l'exigence TFE d'un déploiement sur serveur Linux.")
spacer()
spoken("La troisième couche, c'est les données : PostgreSQL 16 hébergé sur Neon.tech, une solution "
       "serverless avec scaling automatique, hébergée en Europe centrale — donc conforme au RGPD.")
spacer()
spoken("En dehors de ces trois couches, le backend communique avec trois services externes : "
       "Groq pour l'intelligence artificielle du chatbot, ElevenLabs pour la synthèse vocale, "
       "et Stripe pour les paiements en ligne.")
spacer()
spoken("Ce que j'aime dans cette architecture, c'est qu'elle est entièrement automatisée : "
       "chaque push sur GitHub redéploie les trois couches sans aucune intervention manuelle.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SLIDE 5 — TECHNOLOGIES
# ══════════════════════════════════════════════════════════════════════════════
slide_header(5, 'Choix Technologiques', '~1 min 30')

spoken("Chaque choix technologique a une justification précise.")
spacer()
spoken("Spring Boot 3 avec Java 17 : c'est le standard industriel pour les backends Java. "
       "L'auto-configuration réduit le code de configuration, et Spring Data JPA simplifie "
       "toutes les opérations en base de données.")
spacer()
spoken("React 19 avec Vite 7 : Vite build en quelques secondes là où webpack prenait des minutes. "
       "Avec le DOM virtuel de React, l'interface est fluide et réactive.")
spacer()
spoken("PostgreSQL : j'avais besoin des ENUMs natifs pour les statuts de commande, de table, "
       "de paiement — et des triggers avancés pour l'automatisation du stock. MySQL n'offre "
       "pas ce niveau de maturité sur ces deux points.")
spacer()
spoken("Groq plutôt qu'OpenAI : c'est le choix le plus important pour le chatbot vocal. "
       "Groq utilise des processeurs LPU spécialisés qui donnent une latence cinq fois plus "
       "faible qu'OpenAI. Pour un chatbot en temps réel, chaque milliseconde compte.")
spacer()
spoken("ElevenLabs : pour avoir une voix française naturelle, pas robotique. La différence "
       "avec la synthèse vocale standard du navigateur est immédiatement perceptible.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SLIDE 6 — LES 6 RÔLES
# ══════════════════════════════════════════════════════════════════════════════
slide_header(6, 'Les 6 Rôles du Système', '~45 secondes')

spoken("L'application distingue six rôles, chacun avec son propre périmètre fonctionnel "
       "et son interface dédiée.")
spacer()
spoken("L'Administrateur supervise tout : le personnel, le stock, les fournisseurs, "
       "et consulte le tableau de bord analytique.")
spacer()
spoken("Le Caissier gère les encaissements et génère les tickets.")
spacer()
spoken("Le Serveur prend les commandes en salle et gère les tables.")
spacer()
spoken("Le Cuisinier reçoit les commandes en temps réel et met à jour les statuts de préparation.")
spacer()
spoken("Le Client commande en ligne, réserve une table, paie par Stripe, et utilise le chatbot vocal.")
spacer()
spoken("Le Fournisseur consulte ses commandes d'approvisionnement et confirme les livraisons.")
spacer()
spoken("Ce système de rôles est géré par un contrôle d'accès RBAC — chaque endpoint Spring Boot "
       "vérifie le rôle avant d'autoriser l'opération.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SLIDE 7 — TITRE DÉMO
# ══════════════════════════════════════════════════════════════════════════════
slide_header(7, 'Titre — Démonstration', '~15 secondes')

spoken("Place maintenant à la démonstration.")
spacer()
spoken("Je vais vous montrer cinq scénarios clés de l'application, "
       "qui tourne en ce moment en production.")

tip_note("Avant de commencer : vérifiez que l'application est ouverte dans le navigateur. "
         "Ayez deux onglets : un connecté Serveur/Admin, un connecté Cuisinier.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SLIDE 8 — DÉMO ADMIN
# ══════════════════════════════════════════════════════════════════════════════
slide_header(8, 'Démo 1 — Dashboard Admin', '~1 min 30')

action_note("Se connecter en tant qu'Administrateur.")
spacer()
spoken("Je me connecte en tant qu'Administrateur.")
spacer()
action_note("Montrer le tableau de bord : chiffre d'affaires, graphiques, top produits.")
spacer()
spoken("Vous voyez ici le tableau de bord. En temps réel, l'administrateur voit le chiffre "
       "d'affaires du jour, de la semaine, du mois. Les graphiques montrent les tendances "
       "de ventes. Et ici, le top des produits les plus vendus.")
spacer()
action_note("Naviguer vers le catalogue produits.")
spacer()
spoken("Dans le catalogue, chaque produit a un seuil d'alerte de stock. Dès que le stock "
       "descend en dessous de ce seuil, une alerte est générée.")
spacer()
spoken("Ce qui est important ici, c'est que cette alerte n'est pas générée par le code Java. "
       "Elle est générée par un trigger PostgreSQL — directement dans la base de données. "
       "Ça garantit que même si l'application plante, la cohérence des données est préservée.")
spacer()
action_note("Montrer le journal d'audit.")
spacer()
spoken("Le journal d'audit trace toutes les actions : qui a modifié quoi, quand. C'est "
       "une exigence de traçabilité importante pour ce type d'application.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SLIDE 9 — DÉMO COMMANDE
# ══════════════════════════════════════════════════════════════════════════════
slide_header(9, 'Démo 2 — Commande en Temps Réel', '~2 minutes')

tip_note("Deux onglets ouverts : Serveur à gauche, Cuisinier à droite.")
spacer()
spoken("Pour ce scénario, j'ai ouvert deux onglets : un connecté en tant que Serveur, "
       "et un connecté en tant que Cuisinier.")
spacer()
action_note("Montrer les deux onglets côte à côte.")
spacer()
spoken("Je vais créer une commande depuis l'interface Serveur pour la table 3.")
spacer()
action_note("Créer la commande : Table 3, 2x Burger, 1x Frites — valider.")
spacer()
spoken("Je sélectionne deux Burgers et une portion de Frites. Je valide la commande.")
spacer()
action_note("Basculer immédiatement vers l'onglet Cuisinier.")
spacer()
spoken("Regardez ici dans l'interface Cuisinier : la commande vient d'apparaître "
       "instantanément, en moins de cent millisecondes.")
spacer()
spoken("C'est ça la puissance du WebSocket STOMP. Le serveur Spring Boot joue le rôle "
       "de broker de messages. Quand une commande est créée, il diffuse une notification "
       "à tous les clients abonnés au topic /topic/orders. Le cuisinier reçoit la commande "
       "sans avoir besoin de rafraîchir sa page.")
spacer()
action_note("Le cuisinier met à jour le statut : 'En préparation' puis 'Prêt'.")
spacer()
spoken("Le cuisinier peut maintenant marquer la commande En préparation, puis Prête. "
       "Et à ce moment, c'est le serveur qui est notifié. "
       "La communication est bidirectionnelle et instantanée.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SLIDE 10 — DÉMO PAIEMENT
# ══════════════════════════════════════════════════════════════════════════════
slide_header(10, 'Démo 3 — Paiement Stripe & Ticket PDF', '~1 min 30')

action_note("Se connecter en tant que Caissier.")
spacer()
spoken("Je me connecte maintenant en tant que Caissier.")
spacer()
action_note("Sélectionner la commande active — choisir méthode STRIPE.")
spacer()
spoken("Je sélectionne la commande active que nous venons de créer. "
       "Je choisis la méthode de paiement Stripe.")
spacer()
action_note("Taper la carte test : 4242 4242 4242 4242 — n'importe quelle date future — CVV 123.")
spacer()
spoken("Vous voyez ici le formulaire Stripe. En mode test, j'utilise la carte "
       "4242 4242 4242 4242 avec n'importe quelle date future et n'importe quel CVV.")
spacer()
action_note("Confirmer le paiement.")
spacer()
spoken("Je confirme le paiement.")
spacer()
spoken("Ce qui se passe en coulisses : notre backend crée un PaymentIntent sur Stripe. "
       "Le navigateur confirme le paiement directement auprès de Stripe — les données "
       "bancaires ne passent jamais par notre serveur. Stripe nous renvoie une confirmation, "
       "notre backend ferme la commande et enregistre la transaction.")
spacer()
action_note("Montrer le ticket PDF généré et téléchargé.")
spacer()
spoken("Et immédiatement, un ticket PDF est généré. Entièrement côté client, via la "
       "librairie jsPDF. Pas de serveur de rendu PDF — ça allège le backend et accélère "
       "la génération.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SLIDE 11 — DÉMO CHATBOT
# ══════════════════════════════════════════════════════════════════════════════
slide_header(11, 'Démo 4 — Chatbot Vocal Intelligent', '~2 min 30')

tip_note("Vérifiez que le micro est autorisé dans le navigateur AVANT de commencer ce slide.")
spacer()
spoken("C'est la fonctionnalité qui me tient le plus à cœur.")
spacer()
action_note("Aller dans l'espace client — activer le chatbot vocal.")
spacer()
spoken("Je vais dans l'espace client et j'active le chatbot vocal.")
spacer()
spoken("Regardez l'indicateur : je suis en mode LISTENING — le micro est actif.")
spacer()
spoken("Je vais lui poser une question à voix haute.")
spacer()
action_note("Parler clairement au micro : \"Qu'est-ce que vous avez comme plats chauds au menu ?\"")
spacer()
spoken("...")
spacer()
action_note("Laisser la réponse s'afficher et se lire à voix haute — ne pas parler par-dessus.")
spacer()
spoken("Vous voyez trois choses se produire : d'abord la transcription de ma voix en texte "
       "par la Web Speech API du navigateur — aucune clé API requise, c'est natif dans Chrome. "
       "Ensuite, le texte est envoyé à Groq qui génère une réponse avec le modèle LLaMA 3.1 "
       "en moins d'une seconde. Et enfin, la réponse est envoyée à ElevenLabs qui la lit "
       "avec une voix française naturelle.")
spacer()
spoken("La machine à états que vous voyez sur ce slide, c'est ce qui garantit que tout se "
       "passe dans le bon ordre : on ne peut pas être en train d'écouter et de parler en "
       "même temps. C'est ce qui évite que le chatbot s'entende lui-même — le micro est "
       "désactivé pendant la lecture.")
spacer()
spoken("Si ElevenLabs est indisponible, un fallback automatique utilise la synthèse vocale "
       "native du navigateur. L'utilisateur ne voit aucune différence fonctionnelle.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SLIDE 12 — POINTS TECHNIQUES
# ══════════════════════════════════════════════════════════════════════════════
slide_header(12, 'Points Techniques Clés', '~2 minutes')

spoken("Je veux souligner trois points techniques qui font la solidité de cette application.")
spacer()
spoken_bold("Premièrement, les triggers PostgreSQL.")
spacer()
spoken("Quand un serveur valide une commande, les articles sont insérés dans la table "
       "order_items. À ce moment précis, un premier trigger trigger_update_stock se déclenche "
       "et décrémente automatiquement le stock du produit. Si ce nouveau stock passe sous le "
       "seuil d'alerte, un second trigger trigger_stock_alert s'active et insère une ligne "
       "dans la table stock_alerts. Tout ça sans une seule ligne de code Java. C'est la base "
       "de données qui garantit la cohérence — même en cas de panne applicative.")
spacer()
spoken_bold("Deuxièmement, le WebSocket STOMP.")
spacer()
spoken("Le broker Spring Boot gère des topics par rôle. Les cuisiniers s'abonnent à "
       "/topic/orders, les admins à /topic/alerts. SockJS assure la compatibilité avec les "
       "navigateurs qui ne supportent pas nativement WebSocket via un fallback en long-polling.")
spacer()
spoken_bold("Troisièmement, la sécurité.")
spacer()
spoken("Les mots de passe sont hachés avec BCrypt, facteur de coût 10 — aucun mot de passe "
       "n'est stocké en clair. Chaque endpoint est protégé par RBAC. Toutes les communications "
       "sont en HTTPS. Et les requêtes utilisent des PreparedStatement — aucun risque "
       "d'injection SQL.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SLIDE 13 — GITHUB CI/CD
# ══════════════════════════════════════════════════════════════════════════════
slide_header(13, 'GitHub & CI/CD', '~45 secondes')

spoken("Le projet est entièrement versionné sur GitHub, avec un historique de commits "
       "régulier qui témoigne de la progression du développement.")
spacer()
spoken("Ce qui est important, c'est que le déploiement est entièrement automatisé. "
       "Render et Vercel sont tous les deux connectés au dépôt GitHub. Dès qu'un push "
       "est effectué sur la branche main, Render rebuild et redéploie le backend Spring Boot "
       "via Maven, et Vercel rebuilde et redistribue le frontend React sur son CDN.")
spacer()
spoken_bold("De la commande git push à l'application en production, il faut moins de trois minutes.")
spacer()
spoken("Les clés API, les mots de passe de base de données, les identifiants Stripe — "
       "tout ça est configuré comme variables d'environnement dans les dashboards Render "
       "et Vercel. Jamais dans le code. Jamais sur GitHub.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SLIDE 14 — DIFFICULTÉS
# ══════════════════════════════════════════════════════════════════════════════
slide_header(14, 'Difficultés Rencontrées', '~1 min 15')

spoken("Je veux être transparent sur les difficultés réelles que j'ai rencontrées.")
spacer()
spoken_bold("La première, c'est l'effet écho du chatbot vocal.")
spacer()
spoken("Quand ElevenLabs lisait une réponse à voix haute, le microphone captait l'audio, "
       "le transcrivait comme une nouvelle question, et relançait une requête Groq en boucle "
       "infinie. La solution : désactiver le micro pendant toute la durée de la lecture audio, "
       "en m'appuyant sur l'événement onended de l'objet Audio. J'ai aussi ajouté un timer "
       "de sécurité de 15 secondes en cas de silence prolongé, et un timer global de "
       "30 secondes pour réinitialiser le chatbot en cas de blocage d'état.")
spacer()
spoken_bold("La deuxième difficulté, c'est la latence du TTS.")
spacer()
spoken("Le modèle ElevenLabs standard ajoutait 2 à 3 secondes d'attente, ce qui rendait "
       "la conversation peu naturelle. En migrant vers le modèle eleven_turbo_v2_5 et en "
       "découpant les réponses en phrases courtes, j'ai divisé la latence perçue par deux.")
spacer()
spoken_bold("La troisième, plus technique : la taille du diagramme MPD.")
spacer()
spoken("Le diagramme avec ses 14 tables générait une URL encodée PlantUML de plus de "
       "8000 caractères, ce qui déclenchait une erreur 400 du serveur web. "
       "Solution : découper le diagramme en deux parties.")
spacer()
spoken("Ces trois problèmes, je les présente parce qu'ils montrent que le développement "
       "d'une vraie application, c'est avant tout de la résolution de problèmes concrets.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SLIDE 15 — CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
slide_header(15, 'Conclusion & Perspectives', '~1 minute')

spoken("Pour conclure.")
spacer()
spoken("L'application Gestion Snack est complète et déployée en production. Elle couvre "
       "six rôles métier avec vingt-six fonctionnalités documentées sous forme de user stories. "
       "L'API REST est entièrement documentée via Swagger UI. Les notifications temps réel "
       "fonctionnent par WebSocket. Le chatbot vocal combine trois technologies spécialisées. "
       "Et le CI/CD est entièrement automatisé depuis GitHub.")
spacer()
spoken("Ce projet m'a permis de travailler sur des sujets que je n'avais pas encore explorés "
       "en formation : les triggers de base de données, les WebSockets, l'intégration d'APIs "
       "d'intelligence artificielle, et le déploiement cloud en production réelle.")
spacer()
spoken("Pour les évolutions futures, les pistes les plus intéressantes sont une application "
       "mobile React Native pour les serveurs, un système de fidélité client, et un tableau "
       "de bord prédictif basé sur le machine learning pour anticiper les ruptures de stock.")
spacer()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Je vous remercie pour votre attention.  Je suis maintenant disponible pour répondre à vos questions.")
r.font.size = Pt(14)
r.font.bold = True
r.font.color.rgb = NAVY

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  QUESTIONS PROBABLES DU JURY
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_heading('Questions Probables du Jury — Réponses Préparées', level=1)
p.runs[0].font.color.rgb = NAVY
p.runs[0].font.size = Pt(16)
p.paragraph_format.space_before = Pt(10)

section_title("Choix technologiques", TEAL)

qa_question("Pourquoi Spring Boot plutôt que Node.js ou Django ?")
qa_answer(
    "Spring Boot est plus adapté à un backend d'entreprise avec des relations complexes "
    "entre entités. La maturité de l'écosystème Spring — Spring Data JPA, Spring Security, "
    "Spring WebSocket — m'a permis de construire rapidement une architecture solide. "
    "Java 17 avec ses fonctionnalités modernes reste un choix très pertinent en 2025."
)

qa_question("Pourquoi PostgreSQL et pas MySQL ?")
qa_answer(
    "J'avais besoin de deux fonctionnalités absentes de MySQL : les ENUMs natifs pour "
    "les statuts métier, et les triggers avancés pour l'automatisation du stock. "
    "PostgreSQL est aussi plus mature sur les performances et le support du JSONB. "
    "Et Neon.tech propose un hébergement PostgreSQL serverless en Europe — pas MySQL."
)

qa_question("Vous n'utilisez pas JWT — pourquoi ?")
qa_answer(
    "Pour le périmètre de cette application, les sessions serveur gérées par Spring Security "
    "suffisent et sont même plus sûres pour un usage web classique. JWT est utile quand on "
    "a plusieurs serveurs stateless ou une architecture microservices — ce qui n'est pas "
    "le cas ici. Ajouter JWT aurait complexifié l'architecture sans bénéfice réel."
)

qa_question("Pourquoi Groq plutôt qu'OpenAI pour le chatbot ?")
qa_answer(
    "La latence. Groq utilise des processeurs LPU spécialisés dans l'inférence de modèles "
    "de langage — la réponse arrive en moins d'une seconde. OpenAI GPT-3.5 ou GPT-4 "
    "donnent des latences de 2 à 5 secondes. Pour un chatbot vocal, cette différence est "
    "complètement perceptible par l'utilisateur. Et le tier gratuit de Groq est suffisant "
    "pour un projet TFE."
)

section_title("Architecture & Déploiement", TEAL)

qa_question("Pourquoi Render pour le backend et pas Heroku ou AWS ?")
qa_answer(
    "Render offre un déploiement automatique depuis GitHub, un certificat SSL gratuit, "
    "et un conteneur Linux — ce qui répond exactement à l'exigence TFE. Heroku est devenu "
    "payant. AWS est trop complexe pour ce périmètre. Render est le meilleur compromis "
    "simplicité/puissance pour un projet TFE."
)

qa_question("Comment gérez-vous la sécurité des clés API ?")
qa_answer(
    "Toutes les clés sensibles — Groq, ElevenLabs, Stripe, la chaîne de connexion "
    "PostgreSQL — sont des variables d'environnement configurées dans les dashboards "
    "Render et Vercel. Le fichier .env est exclu du dépôt via .gitignore. "
    "Le code source ne contient aucune clé en dur."
)

qa_question("Votre application supporte-t-elle plusieurs snacks simultanément ?")
qa_answer(
    "Non, dans cette version. L'architecture est mono-établissement. C'est d'ailleurs "
    "une des perspectives d'évolution : le support multi-établissements. "
    "Techniquement, cela nécessiterait d'ajouter une table 'establishment' et de "
    "scoper toutes les requêtes par établissement."
)

section_title("Technique & Code", TEAL)

qa_question("Comment fonctionnent exactement les triggers PostgreSQL ?")
qa_answer(
    "Il y en a deux. Le premier, trigger_update_stock, est déclenché AFTER INSERT sur "
    "order_items. Il exécute une fonction qui fait UPDATE products SET stock = stock - quantity "
    "WHERE id = product_id. Le second, trigger_stock_alert, est déclenché AFTER UPDATE sur "
    "products. Si le nouveau stock est inférieur à alert_threshold, il insère une ligne dans "
    "stock_alerts. Ces triggers s'exécutent dans la transaction SQL — ils sont atomiques."
)

qa_question("Comment fonctionne le chatbot quand il n'y a pas de connexion internet ?")
qa_answer(
    "Le chatbot vocal nécessite internet pour Groq et ElevenLabs. Sans connexion, "
    "il ne fonctionne pas. C'est une limitation assumée. La reconnaissance vocale "
    "Web Speech API fonctionne localement sur Chrome, mais la génération de réponse "
    "et la synthèse vocale nécessitent les APIs cloud. Le fallback WebSpeech Synthesis "
    "n'aide que pour la lecture, pas pour la génération."
)

qa_question("Avez-vous écrit des tests unitaires ?")
qa_answer(
    "J'ai écrit des tests pour les couches les plus critiques, notamment les services "
    "de gestion des commandes et du stock. En toute transparence, la couverture de tests "
    "n'est pas aussi complète que je l'aurais souhaité — c'est d'ailleurs quelque chose "
    "que je mettrais en priorité si je refaisais le projet : TDD dès le début."
)

section_title("Questions pièges possibles", RED)

qa_question("Qu'est-ce que vous changeriez si vous refaisiez le projet ?")
qa_answer(
    "Trois choses. D'abord les tests — je les aurais mis en place dès le début en TDD. "
    "Ensuite la documentation API Swagger — je l'aurais maintenue en parallèle du "
    "développement plutôt qu'en fin de projet. Et enfin Docker pour normaliser "
    "l'environnement de développement entre différentes machines."
)

qa_question("Quelle est la partie du projet dont vous êtes le moins satisfait ?")
qa_answer(
    "La couverture de tests. Dans un projet professionnel, on viserait 70 à 80% "
    "de couverture. Ici, j'ai priorisé les fonctionnalités pour avoir une application "
    "complète et démontrable. C'est un compromis conscient, mais c'est clairement "
    "un point d'amélioration important."
)

qa_question("Le chatbot comprend-il vraiment le contexte du snack ?")
qa_answer(
    "Oui, grâce au system prompt. Quand le frontend envoie une requête à Groq, "
    "il inclut un message système qui définit le contexte : le nom du snack, "
    "le menu, les horaires, et les instructions pour que le chatbot réponde "
    "uniquement sur les sujets pertinents. LLaMA 3.1 suit ces instructions "
    "de façon très fiable."
)

# ── Fin ───────────────────────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('— Bonne présentation ! —')
r.font.size  = Pt(14)
r.font.italic = True
r.font.color.rgb = TEAL

output = r"d:\Projet TFE\gestion-snack\fichiers\Script_Oral_TFE_Gestion_Snack.docx"
doc.save(output)
print(f"[OK] Script oral généré : {output}")
