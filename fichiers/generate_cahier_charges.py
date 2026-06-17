# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Marges ────────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Palette de couleurs ───────────────────────────────────────────────────────
C_NAVY  = RGBColor(0x1F, 0x39, 0x6D)
C_TEAL  = RGBColor(0x00, 0x7B, 0x83)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_DARK  = RGBColor(0x1A, 0x1A, 0x2E)
C_GREY  = RGBColor(0x55, 0x55, 0x55)

# ── Fond de cellule ───────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

# ── Titre de niveau 1 ─────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = C_NAVY
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '4')
    bot.set(qn('w:color'), '1F396D')
    pb.append(bot)
    pPr.append(pb)

# ── Titre de niveau 2 ─────────────────────────────────────────────────────────
def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = C_TEAL

# ── Paragraphe courant ────────────────────────────────────────────────────────
def body(text, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = bold
    run.font.color.rgb = C_DARK

# ── Puce avec coche ───────────────────────────────────────────────────────────
def req(text, level=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.8 + level * 0.5)
    r1 = p.add_run('►  ')
    r1.font.color.rgb = C_TEAL
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    r2.font.color.rgb = C_DARK

# ── Note italique ─────────────────────────────────────────────────────────────
def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.italic    = True
    run.font.color.rgb = C_GREY

# ── Tableau à deux colonnes ───────────────────────────────────────────────────
def table2(rows_data, header=None, col1_w=4.5, col2_w=10.5):
    n = len(rows_data) + (1 if header else 0)
    t = doc.add_table(rows=n, cols=2)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.columns[0].width = Cm(col1_w)
    t.columns[1].width = Cm(col2_w)
    r = 0
    if header:
        for ci, h in enumerate(header):
            c = t.rows[r].cells[ci]
            set_cell_bg(c, '1F396D')
            run = c.paragraphs[0].add_run(h)
            run.bold = True; run.font.color.rgb = C_WHITE; run.font.size = Pt(11)
        r += 1
    for row in rows_data:
        for ci, val in enumerate(row):
            c = t.rows[r].cells[ci]
            if r % 2 == 0: set_cell_bg(c, 'EFF6FF')
            run = c.paragraphs[0].add_run(val)
            run.font.size = Pt(10); run.font.color.rgb = C_DARK
        r += 1
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  PAGE DE GARDE
# ══════════════════════════════════════════════════════════════════════════════
for _ in range(5): doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('CAHIER DES CHARGES')
run.bold = True; run.font.size = Pt(28); run.font.color.rgb = C_NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Application Web de Gestion de Snack')
run.bold = True; run.font.size = Pt(18); run.font.color.rgb = C_TEAL

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Travail de Fin d\'Études (TFE)')
run.font.size = Pt(13); run.font.color.rgb = C_GREY

for _ in range(4): doc.add_paragraph()

for label, val in [
    ('Auteur',      'Tiegni Bernard Joël'),
    ('Email',       'tiegnigamobernardjoel@gmail.com'),
    ('Formation',   'Bachelier en Informatique de Gestion'),
    ('Promotion',   '2025 – 2026'),
    ('Version',     'V1 – Document initial'),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f'{label} : ')
    r1.bold = True; r1.font.size = Pt(12); r1.font.color.rgb = C_NAVY
    r2 = p.add_run(val)
    r2.font.size = Pt(12); r2.font.color.rgb = C_DARK

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  1. CONTEXTE ET PROBLÉMATIQUE
# ══════════════════════════════════════════════════════════════════════════════
h1('1. Contexte et problématique')

body(
    'Le snack Tiegni Bernard Joël est un établissement de restauration rapide dont l\'ensemble '
    'des opérations quotidiennes - prise de commandes, encaissement, gestion des stocks, '
    'organisation du personnel - est actuellement géré manuellement ou à l\'aide d\'outils '
    'non intégrés. Cette situation engendre plusieurs difficultés récurrentes :'
)
req('Perte de temps significative lors de la prise et du suivi des commandes')
req('Absence de traçabilité des opérations : impossible de retrouver l\'historique des ventes')
req('Erreurs fréquentes dans la gestion des stocks (ruptures non détectées, sur-commandes)')
req('Aucun outil de réservation accessible en ligne pour les clients')
req('Pas de visibilité en temps réel sur l\'état d\'occupation des tables')
req('Impossibilité de consulter le chiffre d\'affaires de façon instantanée')
req('Aucun canal de communication direct et interactif entre le snack et sa clientèle')

doc.add_paragraph()
body(
    'Face à ces problèmes, le snack souhaite se doter d\'une solution informatisée centralisée, '
    'accessible depuis n\'importe quel appareil connecté et couvrant l\'ensemble des processus '
    'internes de l\'établissement.'
)

# ══════════════════════════════════════════════════════════════════════════════
#  2. OBJECTIFS DU PROJET
# ══════════════════════════════════════════════════════════════════════════════
h1('2. Objectifs du projet')

h2('2.1  Objectifs fonctionnels')
body('L\'application devra permettre de :')
req('Centraliser la gestion des commandes, des stocks, du personnel, des tables et des fournisseurs dans une seule application')
req('Permettre aux clients de consulter le menu, d\'effectuer des réservations et de laisser des avis en ligne')
req('Automatiser la mise à jour des stocks à chaque vente')
req('Gérer les accès via un système de rôles (six profils distincts)')
req('Suivre en temps réel le statut des commandes et l\'occupation des tables')
req('Enregistrer toutes les actions critiques dans un journal d\'audit')
req('Générer des rapports et des exports sur les ventes et le chiffre d\'affaires')
req('Intégrer un assistant virtuel intelligent (chatbot) pour conseiller clients et personnel')
req('Proposer un module de communication vocale entre le client et l\'assistant')
req('Traiter les paiements en ligne par carte bancaire')
req('Notifier en temps réel les utilisateurs concernés lors des événements clés')

h2('2.2  Objectifs non fonctionnels')
req('L\'interface devra être responsive et accessible sur mobile, tablette et desktop')
req('Les temps de réponse de l\'application devront rester acceptables en conditions normales d\'utilisation')
req('L\'application devra être disponible en production de façon continue')
req('L\'accès aux données devra être sécurisé selon les rôles de chaque utilisateur')
req('Le code source devra être versionné et hébergé sur GitHub avec un suivi de l\'historique')
req('L\'application devra être déployée sur un serveur Linux accessible en ligne')

# ══════════════════════════════════════════════════════════════════════════════
#  3. PÉRIMÈTRE DU PROJET
# ══════════════════════════════════════════════════════════════════════════════
h1('3. Périmètre du projet')

h2('3.1  Dans le périmètre (in scope)')
req('Gestion complète des commandes (création, suivi, clôture)')
req('Gestion des tables et réservations')
req('Gestion des produits et des stocks avec alertes automatiques')
req('Gestion des paiements (liquide, carte, paiement en ligne)')
req('Gestion des employés et attribution des rôles')
req('Gestion des fournisseurs et des approvisionnements')
req('Espace client : inscription, réservation, avis, historique')
req('Chatbot textuel et vocal basé sur une IA')
req('Notifications temps réel pour tous les rôles')
req('Tableau de bord administrateur avec indicateurs clés')
req('Journal d\'audit des actions sensibles')
req('Export de rapports et tickets en PDF')
req('Déploiement en production sur infrastructure cloud (serveur Linux)')

h2('3.2  Hors périmètre (out of scope)')
req('Application mobile native (iOS / Android) - l\'application web responsive suffira')
req('Gestion des livraisons à domicile')
req('Module de fidélité / points clients')
req('Intégration avec des logiciels de comptabilité tiers')
req('Fonctionnalités multi-établissements')

# ══════════════════════════════════════════════════════════════════════════════
#  4. ÉTUDE DE L'EXISTANT
# ══════════════════════════════════════════════════════════════════════════════
h1('4. Étude de l\'existant')

body(
    'Avant d\'entamer le développement, une analyse des principales solutions disponibles sur '
    'le marché a été conduite afin de justifier la nécessité d\'une solution sur mesure.'
)

table2(
    header=['Solution existante', 'Limites identifiées'],
    rows_data=[
        ['SumUp POS',
         'Logiciel de caisse simple pour petits commerces. Pas de gestion multi-rôles, '
         'pas d\'assistant virtuel, pas de réservation en ligne, pas de module fournisseurs. Offre payante.'],
        ['Lightspeed Restaurant',
         'Solution complète pour restaurants. Très onéreuse (50 à 200 €/mois), complexe '
         'à configurer, sans chatbot IA ni assistant vocal intégré.'],
        ['Tiller by SumUp',
         'Logiciel de caisse orienté restauration. Pas de module fournisseur, '
         'pas de réservation en ligne, pas d\'intelligence artificielle conversationnelle.'],
        ['Toast POS',
         'Solution américaine. Non disponible en Belgique, interface uniquement anglophone, '
         'coût élevé et non adapté au contexte local.'],
        ['Zenchef / TheFork',
         'Spécialisé uniquement dans la réservation en ligne. '
         'Aucune gestion interne (stocks, commandes, caisse, personnel).'],
    ],
    col1_w=4, col2_w=11
)

body(
    'Aucune solution existante ne répond à l\'ensemble des besoins du snack à coût raisonnable. '
    'La solution à développer se distinguera par :'
)
req('Un système multi-rôles complet (6 profils) dans une seule application intégrée')
req('Un assistant vocal temps réel avec reconnaissance vocale et synthèse TTS')
req('Un chatbot IA disponible en mode texte et vocal, en plusieurs langues')
req('Des notifications temps réel pour tous les acteurs du snack')
req('Un déploiement entièrement sur infrastructure cloud gratuite, sans coût de licence')
req('Un code source ouvert et versionné sur GitHub')

# ══════════════════════════════════════════════════════════════════════════════
#  5. ACTEURS ET RÔLES
# ══════════════════════════════════════════════════════════════════════════════
h1('5. Acteurs et rôles')

body(
    'L\'application devra distinguer six types d\'utilisateurs, chacun disposant d\'un espace '
    'dédié et d\'autorisations propres à ses responsabilités.'
)

table2(
    header=['Rôle', 'Responsabilités attendues'],
    rows_data=[
        ['Administrateur',
         'Gérer les comptes employés et attribuer les rôles. Superviser l\'ensemble du système. '
         'Consulter le chiffre d\'affaires et les indicateurs clés. Gérer les fournisseurs et les produits. '
         'Consulter le journal d\'audit.'],
        ['Caissier',
         'Enregistrer les paiements (liquide, carte, paiement en ligne). '
         'Clôturer les commandes après encaissement. Générer les tickets de caisse. '
         'Consulter l\'historique des transactions.'],
        ['Serveur',
         'Créer et suivre les commandes (sur place ou à emporter). '
         'Gérer l\'occupation des tables. Recevoir les notifications en temps réel.'],
        ['Cuisinier',
         'Consulter la file d\'attente des commandes à préparer. '
         'Mettre à jour le statut de chaque commande. Recevoir les nouvelles commandes en temps réel.'],
        ['Client',
         'S\'inscrire et se connecter (activation par email). Consulter le menu. '
         'Effectuer des réservations de tables. Laisser des avis. '
         'Interagir avec le chatbot textuel et vocal.'],
        ['Fournisseur',
         'Gérer les produits liés à ses approvisionnements. '
         'Consulter et mettre à jour les informations de livraison.'],
    ],
    col1_w=3.5, col2_w=11.5
)

# ══════════════════════════════════════════════════════════════════════════════
#  6. BESOINS FONCTIONNELS
# ══════════════════════════════════════════════════════════════════════════════
h1('6. Besoins fonctionnels')

h2('6.1  Authentification et gestion des comptes')
req('Le système devra permettre la création de comptes utilisateurs avec attribution de rôle')
req('L\'accès devra être conditionné à une authentification par identifiant et mot de passe')
req('Les clients devront activer leur compte via un lien reçu par email avant de pouvoir se connecter')
req('Un utilisateur devra pouvoir réinitialiser son mot de passe par email')
req('Un administrateur devra pouvoir désactiver ou réactiver un compte employé')
req('Chaque utilisateur devra pouvoir modifier les informations de son profil')

h2('6.2  Gestion des commandes')
req('Le serveur devra pouvoir créer une commande en précisant le mode (sur place ou à emporter)')
req('Une commande devra pouvoir contenir plusieurs produits avec des options (extras, quantités)')
req('Le statut d\'une commande devra évoluer selon un cycle défini : reçue → en préparation → prête → payée')
req('Les changements de statut devront être communiqués en temps réel aux acteurs concernés')
req('Une commande devra pouvoir être annulée avant sa clôture')

h2('6.3  Gestion des tables')
req('Le système devra afficher en temps réel le statut de chaque table (disponible, occupée, réservée)')
req('Le serveur devra pouvoir assigner une commande à une table')
req('Le serveur devra pouvoir libérer une table après le départ du client')

h2('6.4  Gestion des produits et des stocks')
req('L\'administrateur devra pouvoir ajouter, modifier et supprimer des produits')
req('Chaque produit devra avoir : nom, prix, catégorie, quantité en stock, seuil d\'alerte')
req('La vente d\'un produit devra automatiquement décrémenter le stock correspondant')
req('Une alerte devra être générée si la quantité en stock passe sous le seuil critique défini')
req('Les alertes de stock devront être consultables et acquittables par l\'administrateur')

h2('6.5  Gestion des paiements')
req('Le caissier devra pouvoir enregistrer un paiement en liquide, par carte ou via paiement en ligne')
req('Un ticket de caisse au format PDF devra être générable pour chaque commande clôturée')
req('L\'historique de toutes les transactions devra être consultable')
req('Le chiffre d\'affaires devra être calculable par période (jour, semaine, mois)')

h2('6.6  Réservations')
req('Un client devra pouvoir réserver une table en précisant la date, l\'heure et le nombre de couverts')
req('Le client devra pouvoir consulter et annuler ses réservations')
req('Le serveur devra pouvoir visualiser les réservations du jour avec les informations du client')

h2('6.7  Avis clients')
req('Un client connecté devra pouvoir laisser un avis avec une note et un commentaire')
req('Les avis devront être visibles publiquement sur la page menu')
req('L\'administrateur devra pouvoir supprimer un avis inapproprié')

h2('6.8  Assistant virtuel (Chatbot)')
req('L\'application devra intégrer un chatbot capable de répondre aux questions des utilisateurs sur le menu et le snack')
req('Le chatbot devra être accessible via une interface textuelle pour tous les rôles')
req('Le chatbot devra proposer un mode vocal : l\'utilisateur parlera au micro, le système répondra vocalement')
req('Le mode vocal devra être compatible avec les principaux navigateurs (Chrome, Safari, Firefox) sur mobile et desktop')
req('Le chatbot vocal devra supporter au minimum le français, le néerlandais et l\'allemand')

h2('6.9  Notifications temps réel')
req('Le système devra envoyer des notifications en temps réel aux utilisateurs concernés lors des événements clés')
req('Les événements notifiés devront inclure au minimum : nouvelle commande, commande prête, annulation, alerte de stock')

h2('6.10  Tableau de bord administrateur')
req('L\'administrateur devra disposer d\'un tableau de bord synthétique affichant les indicateurs clés du snack')
req('Le tableau de bord devra afficher : chiffre d\'affaires du jour, nombre de commandes en cours, alertes en attente')

h2('6.11  Journal d\'audit')
req('Toutes les actions sensibles (connexions, ajouts, suppressions, paiements) devront être enregistrées automatiquement')
req('L\'administrateur devra pouvoir consulter l\'historique complet du journal d\'audit')

h2('6.12  Exports PDF')
req('L\'application devra permettre l\'export en PDF des tickets de caisse et des rapports de ventes')

# ══════════════════════════════════════════════════════════════════════════════
#  7. BESOINS NON FONCTIONNELS
# ══════════════════════════════════════════════════════════════════════════════
h1('7. Besoins non fonctionnels')

table2(
    header=['Critère', 'Exigence'],
    rows_data=[
        ['Performance',      'Les temps de réponse de l\'API devront rester raisonnables en conditions normales d\'utilisation'],
        ['Disponibilité',    'L\'application devra être accessible en ligne de façon continue, avec redéploiement automatique'],
        ['Sécurité',         'Les mots de passe devront être stockés de façon sécurisée (hachage). Les accès devront être contrôlés par rôle. Les communications devront transiter en HTTPS'],
        ['Compatibilité',    'L\'interface devra fonctionner sur les principaux navigateurs modernes (Chrome, Firefox, Safari, Edge) et être responsive'],
        ['Maintenabilité',   'L\'architecture devra être structurée en couches distinctes. Le code devra être versionné sur GitHub avec des messages de commit explicites'],
        ['Évolutivité',      'L\'architecture devra être conçue pour accueillir de futures fonctionnalités sans refonte majeure'],
        ['Déploiement',      'L\'application devra être déployée sur serveur Linux accessible en production'],
    ],
    col1_w=4, col2_w=11
)

# ══════════════════════════════════════════════════════════════════════════════
#  8. CONTRAINTES DU PROJET
# ══════════════════════════════════════════════════════════════════════════════
h1('8. Contraintes du projet')

h2('8.1  Contraintes techniques')
req('Le backend devra être développé en Java avec le framework Spring Boot')
req('Le frontend devra être une application web (SPA) développée en React')
req('La base de données utilisée devra être PostgreSQL')
req('L\'architecture devra respecter la séparation en couches : Présentation / Métier / Accès aux données')
req('Le projet devra obligatoirement être hébergé sur GitHub avec un historique de commits complet')
req('L\'application devra être déployée sur un serveur Linux accessible en ligne')
req('Un Dockerfile devra être fourni pour permettre un déploiement alternatif sur VM')
req('La documentation de l\'API REST devra être générée automatiquement (Swagger / OpenAPI)')

h2('8.2  Contraintes de sécurité')
req('Les mots de passe devront être hachés avec un algorithme robuste (BCrypt)')
req('L\'activation des comptes clients devra passer par une confirmation par email')
req('Toutes les communications entre frontend et backend devront utiliser HTTPS')
req('Les données sensibles (clés API, mots de passe) ne devront jamais être présentes dans le code source')

h2('8.3  Contraintes de budget')
req('L\'infrastructure de déploiement devra avoir un coût nul ou minimal (plans gratuits des plateformes cloud)')
req('Les APIs tierces intégrées devront être disponibles sur des plans gratuits pour le contexte TFE')

h2('8.4  Contraintes de délai')
req('Le projet complet (développement + documentation + déploiement) devra être livré pour la session de présentation TFE 2025-2026')

# ══════════════════════════════════════════════════════════════════════════════
#  9. RÈGLES MÉTIER
# ══════════════════════════════════════════════════════════════════════════════
h1('9. Règles métier')

req('Chaque utilisateur possédera un rôle unique parmi : Administrateur, Caissier, Serveur, Cuisinier, Client, Fournisseur')
req('Un employé désactivé ne pourra plus se connecter au système')
req('Les clients devront confirmer leur email avant leur première connexion')
req('Une commande pourra être de type « sur place » (associée à une table) ou « à emporter »')
req('Une commande pourra contenir un ou plusieurs produits avec des options (extras, sauces, quantités)')
req('La vente d\'un produit déclenchera automatiquement la mise à jour du stock correspondant')
req('Si la quantité d\'un produit passe sous le seuil d\'alerte, une notification sera générée pour l\'administrateur')
req('Un ticket de caisse sera généré après la clôture de chaque paiement')
req('Toutes les opérations sensibles seront enregistrées dans le journal d\'audit avec l\'identité de l\'auteur et l\'horodatage')
req('Une réservation devra obligatoirement préciser : date, heure, nombre de couverts et nom du client')
req('Le chatbot vocal ne prendra en compte la parole de l\'utilisateur que lorsque le système est en état d\'écoute active')

# ══════════════════════════════════════════════════════════════════════════════
#  10. LIVRABLES ATTENDUS
# ══════════════════════════════════════════════════════════════════════════════
h1('10. Livrables attendus')

table2(
    header=['Livrable', 'Description'],
    rows_data=[
        ['Cahier des charges',        'Ce document : besoins, contraintes, acteurs, règles métier'],
        ['Cahier d\'analyse',          'Diagrammes UML (cas d\'utilisation, classes, séquences), MCD, MPD, schéma relationnel, user stories, choix technologiques justifiés'],
        ['Script SQL complet',        'Script d\'initialisation de la base de données (tables, contraintes, triggers, données de départ)'],
        ['Application web déployée',  'Application fonctionnelle accessible en ligne via URL publique'],
        ['Code source sur GitHub',    'Dépôt versionné avec historique complet des commits'],
        ['Documentation technique',   'README couvrant : architecture, installation locale, déploiement en production, endpoints API'],
        ['Documentation API Swagger', 'Interface Swagger accessible en production sur le serveur de déploiement'],
        ['Rapport final TFE',         'Document de synthèse du projet pour la présentation orale'],
        ['Présentation orale',        'Démonstration de l\'application devant le jury'],
    ],
    col1_w=4.5, col2_w=10.5
)

# ══════════════════════════════════════════════════════════════════════════════
#  11. PLANIFICATION PRÉVISIONNELLE
# ══════════════════════════════════════════════════════════════════════════════
h1('11. Planification prévisionnelle')

table2(
    header=['Phase', 'Contenu'],
    rows_data=[
        ['Phase 1 - Analyse et conception',
         'Rédaction du cahier des charges. Étude de l\'existant. '
         'Élaboration des diagrammes (MCD, MPD, diagramme de classes). '
         'Définition des user stories et des maquettes.'],
        ['Phase 2 - Mise en place de l\'infrastructure',
         'Configuration du dépôt GitHub. Initialisation du projet Spring Boot. '
         'Initialisation du projet React. Mise en place de la base de données PostgreSQL. '
         'Premier déploiement sur serveur Linux.'],
        ['Phase 3 - Développement du backend',
         'Création des entités, repositories, services et contrôleurs REST. '
         'Mise en place des triggers SQL. Configuration Swagger. '
         'Gestion des rôles et de la sécurité.'],
        ['Phase 4 - Développement du frontend',
         'Création des composants React par rôle. '
         'Intégration des appels API. Mise en page responsive (Tailwind).'],
        ['Phase 5 - Fonctionnalités avancées',
         'Intégration WebSocket (notifications temps réel). '
         'Intégration chatbot IA. Module vocal. '
         'Paiement en ligne. Export PDF.'],
        ['Phase 6 - Tests, corrections et finalisation',
         'Tests fonctionnels sur navigateurs et appareils mobiles. '
         'Corrections de bugs. Optimisation. Rédaction de la documentation finale.'],
    ],
    col1_w=4.5, col2_w=10.5
)

# ══════════════════════════════════════════════════════════════════════════════
#  12. BUDGET PRÉVISIONNEL
# ══════════════════════════════════════════════════════════════════════════════
h1('12. Budget prévisionnel')

table2(
    header=['Poste de dépense', 'Estimation'],
    rows_data=[
        ['Temps de développement (environ 200 heures)', '~3 000 € (à titre indicatif, 15 €/h)'],
        ['Hébergement backend (serveur Linux cloud)',    '0 € - plan gratuit (Render)'],
        ['Hébergement frontend (CDN)',                   '0 € - plan gratuit (Vercel)'],
        ['Base de données cloud',                        '0 € - plan gratuit (Neon.tech)'],
        ['APIs tierces (IA, TTS)',                       '0 € - plans gratuits disponibles'],
        ['Paiement en ligne (Stripe)',                   '0 € - mode test pour le TFE'],
        ['Nom de domaine personnalisé (optionnel)',       '~10 à 15 €/an'],
        ['TOTAL infrastructure',                         '0 à 15 € (hors temps de développement)'],
    ],
    col1_w=8.5, col2_w=6.5
)
note(
    'Note : en cas de mise en production commerciale réelle, des plans payants seraient nécessaires '
    'pour garantir des performances et une disponibilité accrues.'
)

# ══════════════════════════════════════════════════════════════════════════════
#  13. CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
h1('13. Conclusion')

body(
    'Ce cahier des charges définit l\'ensemble des besoins fonctionnels et non fonctionnels '
    'de l\'application de gestion du snack Tiegni Bernard Joël. Il constitue le contrat de '
    'référence entre la demande du client et le travail de développement à réaliser.'
)
doc.add_paragraph()
body(
    'Le projet à développer couvrira six profils utilisateurs distincts, intégrera des technologies '
    'modernes (IA conversationnelle, synthèse vocale, paiement en ligne, notifications temps réel) '
    'et sera déployé sur une infrastructure cloud Linux accessible publiquement. '
    'Il constitue une démonstration complète des compétences en analyse, développement et déploiement '
    'attendues dans le cadre du Travail de Fin d\'Études du Bachelier en Informatique de Gestion.'
)

# ══════════════════════════════════════════════════════════════════════════════
#  SAUVEGARDE
# ══════════════════════════════════════════════════════════════════════════════
out = r'd:/Projet TFE/gestion-snack/fichiers/Cahier_des_charges_Gestion_Snack_V3_Final.docx'
doc.save(out)
print(f'Saved: {out}')
