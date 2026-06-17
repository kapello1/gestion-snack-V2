# -*- coding: utf-8 -*-
"""
Génération du Cahier d'Analyse — Gestion Snack V3 Final
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Marges ──────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.5)

# ── Couleurs ─────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1A, 0x23, 0x5E)
TEAL   = RGBColor(0x00, 0x7A, 0x8A)
ORANGE = RGBColor(0xE6, 0x7E, 0x22)
LGRAY  = RGBColor(0xF2, 0xF4, 0xF7)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
DKGRAY = RGBColor(0x33, 0x33, 0x33)

# ── Helpers ──────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def h1(text: str):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = NAVY
    p.runs[0].font.size      = Pt(16)
    p.runs[0].bold           = True
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(8)
    return p

def h2(text: str):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = TEAL
    p.runs[0].font.size      = Pt(13)
    p.runs[0].bold           = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p

def h3(text: str):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = ORANGE
    p.runs[0].font.size      = Pt(11)
    p.runs[0].bold           = True
    return p

def body(text: str, bold=False, italic=False):
    p = doc.add_paragraph(text)
    for r in p.runs:
        r.font.size  = Pt(10.5)
        r.font.color.rgb = DKGRAY
        r.bold   = bold
        r.italic = italic
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(text: str, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    for r in p.runs:
        r.font.size = Pt(10.5)
        r.font.color.rgb = DKGRAY
    p.paragraph_format.space_after = Pt(2)
    return p

def numbered(text: str, level=0):
    p = doc.add_paragraph(text, style='List Number')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    for r in p.runs:
        r.font.size = Pt(10.5)
        r.font.color.rgb = DKGRAY
    p.paragraph_format.space_after = Pt(2)
    return p

def code_block(text: str):
    """Bloc de code avec fond gris clair et police monospace."""
    for line in text.split('\n'):
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Inches(0.3)
        # fond gris
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  'F0F0F0')
        pPr.append(shd)

def table_header(tbl, headers: list, bg='1A235E'):
    row = tbl.rows[0]
    for i, h in enumerate(headers):
        c = row.cells[i]
        c.text = h
        set_cell_bg(c, bg)
        for p in c.paragraphs:
            for r in p.runs:
                r.font.bold  = True
                r.font.color.rgb = WHITE
                r.font.size  = Pt(9.5)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_row(tbl, values: list, bg=None):
    row = tbl.add_row()
    for i, v in enumerate(values):
        c = row.cells[i]
        c.text = str(v)
        if bg:
            set_cell_bg(c, bg)
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9.5)
                r.font.color.rgb = DKGRAY

def hr():
    doc.add_paragraph('─' * 80)

def page_break():
    doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  PAGE DE GARDE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph('\n\n')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('CAHIER D\'ANALYSE')
r.font.size  = Pt(28)
r.font.bold  = True
r.font.color.rgb = NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Projet TFE — Gestion Snack')
r.font.size  = Pt(18)
r.font.color.rgb = TEAL

doc.add_paragraph('\n')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Application Web de Gestion d\'un Snack\n(Commandes • Stock • Caisse • Réservations • Chatbot Vocal)')
r.font.size  = Pt(12)
r.font.color.rgb = DKGRAY
r.italic = True

doc.add_paragraph('\n\n')
info_tbl = doc.add_table(rows=5, cols=2)
info_tbl.style = 'Table Grid'
pairs = [
    ('Auteur',      'Tiegni Gamo Bernard-Joël'),
    ('Établissement','Institut Libre Marie Haps — Section Informatique'),
    ('Année',       '2025–2026'),
    ('Version',     'V3.0 — Finale'),
    ('Date',        '16 juin 2026'),
]
for i, (k, v) in enumerate(pairs):
    row = info_tbl.rows[i]
    row.cells[0].text = k
    row.cells[1].text = v
    set_cell_bg(row.cells[0], '1A235E')
    for p2 in row.cells[0].paragraphs:
        for r2 in p2.runs:
            r2.font.color.rgb = WHITE
            r2.font.bold = True
            r2.font.size = Pt(10)
    for p2 in row.cells[1].paragraphs:
        for r2 in p2.runs:
            r2.font.size = Pt(10)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
h1('1. Introduction')
body(
    "Ce cahier d'analyse constitue le document technique de référence du projet TFE "
    "«Gestion Snack». Il traduit les exigences fonctionnelles du cahier des charges en "
    "une architecture concrète, détaillant les acteurs, les cas d'utilisation, les "
    "user stories, les diagrammes UML, le modèle de données, le dictionnaire de données "
    "et les choix technologiques justifiés. Ce document respecte la consigne TFE et le "
    "dossier pédagogique de l'établissement."
)

h2('1.1 Objectif du document')
bullet("Décrire l'ensemble des fonctionnalités à développer sous forme de user stories.")
bullet("Fournir les diagrammes UML complets (cas d'utilisation, classes, séquences).")
bullet("Définir le modèle conceptuel et physique de données.")
bullet("Justifier les choix technologiques retenus.")
bullet("Servir de base contractuelle entre l'analyse et le développement.")

h2('1.2 Périmètre')
body(
    "L'application «Gestion Snack» est une plateforme web full-stack couvrant la gestion "
    "opérationnelle complète d'un établissement de restauration rapide : personnel, "
    "commandes, stock, caisse, réservations, chatbot vocal intelligent et tableau de bord "
    "analytique en temps réel."
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  2. ÉTUDE DE L'EXISTANT
# ══════════════════════════════════════════════════════════════════════════════
h1('2. Étude de l\'existant')
body(
    "Avant de concevoir le système, une analyse des solutions existantes sur le marché "
    "a été réalisée afin d'identifier les lacunes et de définir la valeur ajoutée du projet."
)

h2('2.1 Solutions analysées')
tbl = doc.add_table(rows=1, cols=4)
tbl.style = 'Table Grid'
table_header(tbl, ['Solution', 'Points forts', 'Limites identifiées', 'Différenciation'])
rows_data = [
    ('Square POS', 'Interface intuitive, paiements intégrés', 'Pas de chatbot, pas de gestion stock avancée, coût élevé', 'Chatbot vocal IA, gestion stock avec alertes automatiques'),
    ('Lightspeed Restaurant', 'Gestion complète, multi-sites', 'Très complexe, inaccessible aux PME, pas de personnalisation', 'Solution légère, personnalisable, open-source friendly'),
    ('Tiller Systems', 'Interface tactile, rapports', 'Pas de rôles multiples granulaires, pas d\'API publique', 'Système de rôles fin (6 rôles), API REST documentée Swagger'),
    ('Application maison simple', 'Spécifique au contexte', 'Monolithique, pas de temps réel, pas de chatbot', 'Architecture moderne, WebSocket temps réel, IA intégrée'),
]
for vals in rows_data:
    add_row(tbl, vals)

doc.add_paragraph()
h2('2.2 Valeur ajoutée du projet')
bullet("Chatbot vocal intelligent combinant la reconnaissance vocale (Web Speech API), la génération de réponses (Groq/LLaMA 3.1) et la synthèse vocale naturelle (ElevenLabs TTS).")
bullet("Gestion des stocks avec décrément automatique via triggers PostgreSQL et alertes en temps réel par WebSocket.")
bullet("Tableau de bord analytique avec statistiques de revenus, tendances et indicateurs clés.")
bullet("Système de réservations de tables en ligne avec confirmation instantanée.")
bullet("Paiements en ligne intégrés via Stripe.")
bullet("Architecture cloud moderne : Render (backend) + Vercel (frontend) + Neon.tech (PostgreSQL serverless).")
bullet("API REST entièrement documentée via SpringDoc OpenAPI / Swagger UI.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  3. ACTEURS ET RÔLES
# ══════════════════════════════════════════════════════════════════════════════
h1('3. Acteurs et Rôles du Système')
body(
    "Le système distingue six rôles humains et un acteur système (traitements automatiques). "
    "Chaque rôle dispose d'un périmètre fonctionnel défini et d'une interface adaptée."
)

tbl = doc.add_table(rows=1, cols=4)
tbl.style = 'Table Grid'
table_header(tbl, ['Rôle (ENUM)', 'Nom métier', 'Description', 'Interface principale'])
roles_data = [
    ('ADMIN', 'Administrateur', 'Supervise l\'ensemble du système : personnel, stock, fournisseurs, tableau de bord, audit', 'Dashboard Admin complet'),
    ('CASHIER', 'Caissier', 'Gère les paiements, clôture les commandes, imprime les tickets de caisse', 'Interface Caisse'),
    ('WAITER', 'Serveur', 'Prend les commandes en salle, gère les tables et les réservations', 'Interface Commandes'),
    ('COOK', 'Cuisinier', 'Reçoit les commandes en temps réel, met à jour les statuts de préparation', 'Interface Cuisine'),
    ('CUSTOMER', 'Client', 'Consulte le menu, réserve une table, passe des commandes en ligne, utilise le chatbot', 'Interface Client / Chatbot'),
    ('PROVIDER', 'Fournisseur', 'Accède à ses commandes d\'approvisionnement, met à jour les livraisons', 'Interface Fournisseur'),
    ('SYSTÈME', 'Système automatique', 'Décrémente le stock, génère les alertes de stock, notifie via WebSocket', 'Triggers PostgreSQL + WebSocket'),
]
for vals in roles_data:
    add_row(tbl, vals)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  4. USER STORIES
# ══════════════════════════════════════════════════════════════════════════════
h1('4. User Stories')
body(
    "Les user stories sont rédigées selon le format standard : "
    "«En tant que [rôle], je veux [action] afin de [bénéfice].» "
    "Chaque user story est associée à des critères d'acceptation mesurables."
)

# ── 4.1 Admin ────────────────────────────────────────────────────────────────
h2('4.1 Administrateur (ADMIN)')

stories_admin = [
    ('US-A01', 'Gérer les employés',
     'En tant qu\'administrateur, je veux créer, modifier et désactiver des comptes employés afin de contrôler l\'accès au système.',
     ['Le formulaire exige nom, email, rôle et mot de passe.',
      'Un email unique est vérifié avant création.',
      'La désactivation empêche la connexion sans supprimer les données.',
      'Un audit log enregistre chaque modification.']),
    ('US-A02', 'Gérer les fournisseurs',
     'En tant qu\'administrateur, je veux gérer le répertoire des fournisseurs afin de suivre les approvisionnements.',
     ['CRUD complet sur nom, contact, email, téléphone.',
      'Impossible de supprimer un fournisseur lié à des produits.']),
    ('US-A03', 'Gérer le catalogue de produits',
     'En tant qu\'administrateur, je veux créer et modifier les produits (FOOD/DRINK/INGREDIENT) afin de maintenir le catalogue à jour.',
     ['Chaque produit a un nom, prix, type, seuil d\'alerte stock.',
      'La modification du prix est tracée dans l\'audit log.',
      'Un produit désactivé n\'apparaît plus dans les commandes.']),
    ('US-A04', 'Consulter le tableau de bord analytique',
     'En tant qu\'administrateur, je veux voir les KPIs (revenus, commandes, top produits) afin de piloter l\'activité.',
     ['Chiffre d\'affaires du jour/semaine/mois affiché.',
      'Graphiques des tendances de ventes.',
      'Top 5 produits les plus vendus identifiés.']),
    ('US-A05', 'Recevoir les alertes de stock',
     'En tant qu\'administrateur, je veux être notifié en temps réel quand un produit passe sous son seuil d\'alerte afin d\'anticiper les ruptures.',
     ['Notification WebSocket instantanée.',
      'Badge visuel sur le tableau de bord.',
      'Historique des alertes consultable.']),
    ('US-A06', 'Consulter les journaux d\'audit',
     'En tant qu\'administrateur, je veux consulter l\'historique des actions utilisateurs afin de garantir la traçabilité.',
     ['Toutes les actions CRUD sont enregistrées avec horodatage et auteur.',
      'Filtrage par utilisateur, date et type d\'action disponible.']),
    ('US-A07', 'Gérer les transactions financières',
     'En tant qu\'administrateur, je veux consulter et exporter les transactions afin de produire des rapports financiers.',
     ['Liste complète des paiements (CASH/CARD/STRIPE).',
      'Export PDF possible.',
      'Filtrage par date et méthode de paiement.']),
]

for us_id, title, story, criteria in stories_admin:
    h3(f'{us_id} — {title}')
    body(story)
    body('Critères d\'acceptation :', bold=True)
    for c in criteria:
        bullet(c)

# ── 4.2 Caissier ─────────────────────────────────────────────────────────────
h2('4.2 Caissier (CASHIER)')

stories_caissier = [
    ('US-C01', 'Traiter un paiement',
     'En tant que caissier, je veux encaisser une commande (espèces, carte ou Stripe) afin de clôturer la transaction.',
     ['Sélection de la méthode de paiement : CASH, CARD, STRIPE.',
      'Pour Stripe : redirection vers le formulaire de paiement sécurisé.',
      'Le statut de la commande passe à CLOSED après paiement confirmé.',
      'Une transaction est créée en base avec le montant et la méthode.']),
    ('US-C02', 'Générer un ticket de caisse',
     'En tant que caissier, je veux générer et imprimer un ticket PDF après chaque paiement afin de remettre un justificatif au client.',
     ['Le ticket contient : date, heure, liste des articles, total, méthode de paiement.',
      'Généré en PDF via jsPDF.',
      'Téléchargeable et imprimable depuis le navigateur.']),
    ('US-C03', 'Consulter les commandes actives',
     'En tant que caissier, je veux voir toutes les commandes ACTIVE afin de savoir lesquelles sont prêtes à être encaissées.',
     ['Liste triée par heure de création.',
      'Indicateur visuel de l\'état (ACTIVE/CLOSED/CANCELLED).',
      'Mise à jour en temps réel via WebSocket.']),
    ('US-C04', 'Annuler une commande',
     'En tant que caissier, je veux annuler une commande erronée afin de corriger les erreurs de saisie.',
     ['Seule une commande ACTIVE peut être annulée.',
      'L\'annulation enregistre le motif et l\'auteur dans l\'audit log.',
      'Le stock est restitué après annulation.']),
]
for us_id, title, story, criteria in stories_caissier:
    h3(f'{us_id} — {title}')
    body(story)
    body('Critères d\'acceptation :', bold=True)
    for c in criteria:
        bullet(c)

# ── 4.3 Serveur ──────────────────────────────────────────────────────────────
h2('4.3 Serveur (WAITER)')

stories_serveur = [
    ('US-W01', 'Créer une commande en salle',
     'En tant que serveur, je veux créer une commande pour une table ou en emporter afin d\'enregistrer les choix du client.',
     ['Sélection du type : ON_SITE (avec numéro de table) ou TAKEAWAY.',
      'Ajout des articles depuis le catalogue en précisant les quantités.',
      'La commande est créée en statut ACTIVE.',
      'Notification automatique envoyée au cuisinier via WebSocket.']),
    ('US-W02', 'Modifier une commande active',
     'En tant que serveur, je veux ajouter ou retirer des articles d\'une commande ACTIVE afin de répondre aux changements du client.',
     ['Seules les commandes ACTIVE sont modifiables.',
      'Les modifications sont tracées dans l\'audit log.',
      'Le cuisinier reçoit une notification de mise à jour.']),
    ('US-W03', 'Gérer l\'état des tables',
     'En tant que serveur, je veux visualiser et mettre à jour le statut des tables (FREE/OCCUPIED/RESERVED) afin d\'optimiser le placement.',
     ['Vue plan de salle avec code couleur par statut.',
      'Changement de statut manuel possible.',
      'Mise à jour automatique lors de la création/clôture d\'une commande.']),
    ('US-W04', 'Gérer les réservations',
     'En tant que serveur, je veux consulter et confirmer les réservations afin de préparer l\'accueil des clients.',
     ['Liste des réservations du jour visible.',
      'Possibilité de confirmer ou annuler une réservation.',
      'Table réservée passe automatiquement en statut RESERVED.']),
]
for us_id, title, story, criteria in stories_serveur:
    h3(f'{us_id} — {title}')
    body(story)
    body('Critères d\'acceptation :', bold=True)
    for c in criteria:
        bullet(c)

# ── 4.4 Cuisinier ─────────────────────────────────────────────────────────────
h2('4.4 Cuisinier (COOK)')

stories_cook = [
    ('US-K01', 'Visualiser les commandes en attente',
     'En tant que cuisinier, je veux voir en temps réel les nouvelles commandes afin de démarrer la préparation sans délai.',
     ['Nouvelles commandes affichées instantanément via WebSocket.',
      'Détail des articles avec quantités clairement visible.',
      'Tri par ordre d\'arrivée.']),
    ('US-K02', 'Mettre à jour le statut de préparation',
     'En tant que cuisinier, je veux indiquer quand une commande est prête afin d\'en informer le serveur.',
     ['Bouton "En préparation" et "Prêt" disponibles.',
      'Le serveur reçoit une notification WebSocket quand la commande est prête.',
      'L\'horodatage de chaque changement de statut est enregistré.']),
    ('US-K03', 'Consulter les alertes de stock',
     'En tant que cuisinier, je veux être alerté quand un ingrédient est insuffisant afin d\'adapter les préparations.',
     ['Alerte visible dans l\'interface cuisine.',
      'Liste des ingrédients en rupture ou sous le seuil d\'alerte.',
      'Mise à jour en temps réel.']),
]
for us_id, title, story, criteria in stories_cook:
    h3(f'{us_id} — {title}')
    body(story)
    body('Critères d\'acceptation :', bold=True)
    for c in criteria:
        bullet(c)

# ── 4.5 Client ───────────────────────────────────────────────────────────────
h2('4.5 Client (CUSTOMER)')

stories_client = [
    ('US-CL01', 'Créer un compte client',
     'En tant que client, je veux m\'inscrire avec mon email et mot de passe afin d\'accéder aux fonctionnalités en ligne.',
     ['Email unique requis.',
      'Mot de passe haché avec BCrypt.',
      'Confirmation de création par message visuel.']),
    ('US-CL02', 'Consulter le menu en ligne',
     'En tant que client, je veux parcourir le menu par catégorie afin de choisir mes articles avant de commander.',
     ['Menu organisé par catégories (FOOD, DRINK).',
      'Photo, description et prix affichés.',
      'Disponibilité en temps réel (hors-stock grisé).']),
    ('US-CL03', 'Passer une commande en ligne',
     'En tant que client, je veux commander en ligne depuis mon téléphone afin de gagner du temps.',
     ['Panier avec ajout/suppression d\'articles.',
      'Choix du type : sur place (avec numéro de table) ou emporter.',
      'Confirmation de commande avec récapitulatif.']),
    ('US-CL04', 'Réserver une table',
     'En tant que client, je veux réserver une table en ligne afin d\'assurer ma place lors de mon arrivée.',
     ['Sélection de la date, heure et nombre de couverts.',
      'Vérification de la disponibilité en temps réel.',
      'Confirmation de réservation affichée.']),
    ('US-CL05', 'Payer en ligne via Stripe',
     'En tant que client, je veux payer ma commande en ligne par carte bancaire afin d\'éviter l\'attente à la caisse.',
     ['Redirection vers le formulaire Stripe sécurisé.',
      'Confirmation de paiement après succès.',
      'Le statut de commande passe à CLOSED automatiquement.']),
    ('US-CL06', 'Utiliser le chatbot vocal',
     'En tant que client, je veux poser des questions vocalement au chatbot afin d\'obtenir des informations sur le menu et les horaires.',
     ['Activation par microphone (Web Speech API).',
      'Réponse textuelle générée par Groq/LLaMA 3.1.',
      'Réponse lue à voix haute par ElevenLabs TTS (voix française naturelle).',
      'Fallback sur la synthèse vocale du navigateur si ElevenLabs indisponible.']),
    ('US-CL07', 'Laisser un avis',
     'En tant que client, je veux laisser un avis sur ma commande afin de partager mon expérience.',
     ['Note de 1 à 5 étoiles.',
      'Commentaire textuel facultatif.',
      'Un avis par commande clôturée.']),
    ('US-CL08', 'Consulter l\'historique de commandes',
     'En tant que client, je veux voir mes commandes passées afin de suivre mes consommations.',
     ['Liste des commandes avec date, total et statut.',
      'Détail des articles par commande consultable.']),
]
for us_id, title, story, criteria in stories_client:
    h3(f'{us_id} — {title}')
    body(story)
    body('Critères d\'acceptation :', bold=True)
    for c in criteria:
        bullet(c)

# ── 4.6 Fournisseur ──────────────────────────────────────────────────────────
h2('4.6 Fournisseur (PROVIDER)')

stories_provider = [
    ('US-P01', 'Accéder à l\'espace fournisseur',
     'En tant que fournisseur, je veux me connecter avec mes identifiants afin d\'accéder à mes commandes d\'approvisionnement.',
     ['Authentification avec email et mot de passe.',
      'Accès limité aux données le concernant uniquement.']),
    ('US-P02', 'Consulter les commandes d\'approvisionnement',
     'En tant que fournisseur, je veux voir les commandes passées par l\'établissement afin de préparer les livraisons.',
     ['Liste des commandes avec articles, quantités et dates.',
      'Statut de chaque commande clairement affiché.']),
    ('US-P03', 'Mettre à jour le statut de livraison',
     'En tant que fournisseur, je veux confirmer une livraison afin de tenir à jour les stocks du snack.',
     ['Bouton de confirmation de livraison disponible.',
      'Le stock est mis à jour automatiquement après confirmation.',
      'L\'administrateur reçoit une notification de livraison.']),
    ('US-P04', 'Consulter le catalogue des produits à approvisionner',
     'En tant que fournisseur, je veux consulter les produits en rupture ou sous le seuil d\'alerte afin de proposer des réapprovisionnements.',
     ['Vue filtrée sur les produits en alerte.',
      'Quantités actuelles et seuils visibles.']),
]
for us_id, title, story, criteria in stories_provider:
    h3(f'{us_id} — {title}')
    body(story)
    body('Critères d\'acceptation :', bold=True)
    for c in criteria:
        bullet(c)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  5. DIAGRAMME DE CAS D'UTILISATION
# ══════════════════════════════════════════════════════════════════════════════
h1('5. Diagramme de Cas d\'Utilisation')
body(
    "Le diagramme suivant représente l'ensemble des cas d'utilisation du système "
    "pour les six rôles humains et le système automatique. "
    "Insérer le code PlantUML ci-dessous dans un éditeur PlantUML (plantuml.com) "
    "pour générer le diagramme."
)

plantuml_uc = '''@startuml Diagramme_UC_Gestion_Snack
left to right direction
skinparam packageStyle rectangle
skinparam actorStyle awesome
skinparam usecase {
  BackgroundColor #EEF6FF
  BorderColor #007A8A
  ArrowColor #1A235E
}

actor "Administrateur"  as ADM
actor "Caissier"        as CSH
actor "Serveur"         as WTR
actor "Cuisinier"       as COK
actor "Client"          as CLT
actor "Fournisseur"     as PRV
actor "Système auto"    as SYS

rectangle "Gestion Snack" {

  package "Authentification" {
    usecase "Se connecter"                as UC_AUTH
    usecase "Se déconnecter"              as UC_LOGOUT
  }

  package "Administration" {
    usecase "Gérer les employés"          as UC_EMP
    usecase "Gérer les fournisseurs"      as UC_FOUR
    usecase "Gérer le catalogue"          as UC_PROD
    usecase "Consulter tableau de bord"   as UC_DASH
    usecase "Consulter revenus / CA"      as UC_CA
    usecase "Consulter audit log"         as UC_AUDIT
    usecase "Consulter alertes stock"     as UC_ALERT
  }

  package "Caisse" {
    usecase "Traiter un paiement"         as UC_PAY
    usecase "Paiement Stripe"             as UC_STRIPE
    usecase "Générer ticket PDF"          as UC_TICKET
    usecase "Consulter transactions"      as UC_TRANS
    usecase "Annuler une commande"        as UC_CANCEL
  }

  package "Salle & Commandes" {
    usecase "Créer une commande"          as UC_CMD
    usecase "Modifier une commande"       as UC_EDIT_CMD
    usecase "Gérer l'état des tables"     as UC_TABLE
    usecase "Gérer les réservations"      as UC_RES_W
  }

  package "Cuisine" {
    usecase "Voir commandes en attente"   as UC_PREP
    usecase "Mettre à jour préparation"   as UC_STATUS
    usecase "Voir alertes stock cuisine"  as UC_STOCK_C
  }

  package "Statistiques" {
    usecase "Voir statistiques ventes"    as UC_STAT
  }

  package "Espace Client" {
    usecase "Créer un compte"             as UC_REGISTER
    usecase "Consulter le menu"           as UC_MENU
    usecase "Commander en ligne"          as UC_ORDER_CL
    usecase "Réserver une table"          as UC_RES
    usecase "Payer en ligne (Stripe)"     as UC_PAY_CL
    usecase "Laisser un avis"             as UC_AVIS
    usecase "Historique commandes"        as UC_HIST
    usecase "Utiliser le chatbot vocal"   as UC_CHAT
  }

  package "Espace Fournisseur" {
    usecase "Voir commandes appro."       as UC_APPRO
    usecase "Confirmer livraison"         as UC_LIVR
    usecase "Voir produits en alerte"     as UC_PROD_PRV
  }

  package "Automatique" {
    usecase "Décrémenter stock auto"      as UC_STOCK_AUTO
    usecase "Générer alerte stock"        as UC_ALERT_AUTO
    usecase "Notifier via WebSocket"      as UC_NOTIF
  }
}

' Authentification — tous les rôles humains
ADM --> UC_AUTH
CSH --> UC_AUTH
WTR --> UC_AUTH
COK --> UC_AUTH
CLT --> UC_AUTH
PRV --> UC_AUTH
ADM --> UC_LOGOUT
CSH --> UC_LOGOUT
WTR --> UC_LOGOUT
COK --> UC_LOGOUT
CLT --> UC_LOGOUT
PRV --> UC_LOGOUT

' Admin
ADM --> UC_EMP
ADM --> UC_FOUR
ADM --> UC_PROD
ADM --> UC_DASH
ADM --> UC_CA
ADM --> UC_AUDIT
ADM --> UC_ALERT
ADM --> UC_STAT
ADM --> UC_TRANS

' Caissier
CSH --> UC_PAY
CSH --> UC_TICKET
CSH --> UC_TRANS
CSH --> UC_CANCEL
UC_PAY ..> UC_STRIPE : <<extend>>

' Serveur
WTR --> UC_CMD
WTR --> UC_EDIT_CMD
WTR --> UC_TABLE
WTR --> UC_RES_W

' Cuisinier
COK --> UC_PREP
COK --> UC_STATUS
COK --> UC_STOCK_C

' Client
CLT --> UC_REGISTER
CLT --> UC_MENU
CLT --> UC_ORDER_CL
CLT --> UC_RES
CLT --> UC_PAY_CL
CLT --> UC_AVIS
CLT --> UC_HIST
CLT --> UC_CHAT

' Fournisseur
PRV --> UC_APPRO
PRV --> UC_LIVR
PRV --> UC_PROD_PRV

' Système automatique
SYS --> UC_STOCK_AUTO
SYS --> UC_ALERT_AUTO
SYS --> UC_NOTIF
UC_STOCK_AUTO ..> UC_ALERT_AUTO : <<include>>
UC_ALERT_AUTO  ..> UC_NOTIF      : <<include>>

@enduml'''

h2('5.1 Code PlantUML — Diagramme UC')
code_block(plantuml_uc)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  6. ARCHITECTURE TECHNIQUE
# ══════════════════════════════════════════════════════════════════════════════
h1('6. Architecture Technique')

h2('6.1 Architecture globale')
body(
    "L'application suit une architecture client-serveur en trois couches déployée sur le cloud :"
)
bullet("Couche présentation : React 19 + Vite 7 + Tailwind CSS 3 — hébergé sur Vercel (CDN mondial).")
bullet("Couche métier : Spring Boot 3 (Java 17) — hébergé sur Render (conteneur Linux).")
bullet("Couche données : PostgreSQL 16 — hébergé sur Neon.tech (serverless, EU-Central Frankfurt).")

h2('6.2 Schéma d\'architecture')
body("[ Client navigateur ] ←HTTPS→ [ Vercel CDN (React/Vite) ] ←REST/WebSocket→ [ Render (Spring Boot) ] ←JDBC→ [ Neon.tech PostgreSQL ]", italic=True)
body("                                                                              ↕ REST                    ↕ REST")
body("                                                                         [ Groq API ]           [ ElevenLabs API ]", italic=True)
body("                                                                         [ Stripe API ]", italic=True)

h2('6.3 Communication temps réel')
body(
    "Les notifications en temps réel (nouvelles commandes, alertes stock) sont transmises "
    "via WebSocket avec le protocole STOMP sur SockJS. Le backend Spring Boot joue le rôle "
    "de broker de messages. Les clients React s'abonnent aux topics correspondants à leur rôle."
)

h2('6.4 Flux de déploiement CI/CD')
bullet("Chaque push sur la branche main de GitHub déclenche un déploiement automatique.")
bullet("Render redéploie le backend Spring Boot (build Maven + démarrage du jar).")
bullet("Vercel redéploie le frontend React (build Vite + CDN distribution).")
bullet("Les variables d\'environnement sensibles (clés API, URL BD) sont gérées dans les dashboards Render/Vercel.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  7. CHOIX TECHNOLOGIQUES
# ══════════════════════════════════════════════════════════════════════════════
h1('7. Choix Technologiques et Justifications')

tbl = doc.add_table(rows=1, cols=4)
tbl.style = 'Table Grid'
table_header(tbl, ['Technologie', 'Version', 'Rôle', 'Justification'])
tech_rows = [
    ('Spring Boot', '3.5.7', 'Backend REST + WebSocket', 'Standard industriel Java, auto-configuration, large écosystème, Spring Data JPA pour l\'ORM'),
    ('Java', '17 LTS', 'Langage backend', 'Version LTS longue durée, records, sealed classes, performances améliorées'),
    ('React', '19', 'Frontend SPA', 'Composants réutilisables, hooks, performance avec le DOM virtuel, large communauté'),
    ('Vite', '7', 'Build tool frontend', 'Build ultra-rapide par rapport à webpack, HMR natif, optimisé pour React'),
    ('Tailwind CSS', '3', 'Styles frontend', 'CSS utilitaire, pas de fichiers CSS séparés, design system cohérent, responsive natif'),
    ('PostgreSQL', '16', 'Base de données', 'SGBD relationnel robuste, ENUM natifs, triggers avancés, support JSONB'),
    ('Neon.tech', 'Serverless', 'Hébergement BD', 'PostgreSQL serverless, scaling automatique, tier gratuit généreux, EU Frankfurt (RGPD)'),
    ('Render', 'Cloud', 'Hébergement backend', 'Déploiement automatique depuis GitHub, certificat SSL gratuit, logs en temps réel'),
    ('Vercel', 'CDN', 'Hébergement frontend', 'CDN mondial, déploiement instantané, edge network, optimisé pour React/Vite'),
    ('Spring Data JPA', '3.x', 'ORM', 'Abstraction Hibernate, requêtes JPQL/Criteria, repositories automatiques'),
    ('SpringDoc OpenAPI', '2.8.9', 'Documentation API', 'Génération automatique de la doc Swagger UI depuis les annotations Java'),
    ('BCrypt', 'via Spring', 'Sécurité mots de passe', 'Algorithme de hachage adaptatif, facteur de coût configurable, standard de l\'industrie'),
    ('WebSocket STOMP', 'via Spring', 'Temps réel', 'Protocole de messagerie sur WebSocket, intégration native Spring, brokers topics'),
    ('SockJS', '1.x', 'Fallback WebSocket', 'Compatibilité navigateurs sans WebSocket natif via long-polling'),
    ('Groq API', 'cloud', 'IA chatbot', 'Inférence ultra-rapide (LPU), modèle LLaMA 3.1-8b-instant, latence < 1s, tier gratuit'),
    ('ElevenLabs API', 'cloud', 'Synthèse vocale TTS', 'Voix française naturelle, modèle eleven_turbo_v2_5, faible latence, API REST simple'),
    ('Web Speech API', 'navigateur', 'Reconnaissance vocale STT', 'API native navigateur (Chrome), pas d\'API key requise, reconnaissance en français'),
    ('Stripe', 'v3', 'Paiements en ligne', 'Standard industrie, 3D Secure natif, webhook events, mode test complet'),
    ('jsPDF + AutoTable', '2.x + 3.x', 'Export PDF', 'Génération PDF côté client React, tableaux formatés, pas de dépendance backend'),
    ('Axios', '1.x', 'Client HTTP', 'Intercepteurs pour auth, gestion d\'erreurs centralisée, plus ergonomique que fetch'),
    ('GitHub', 'cloud', 'Versioning + CI/CD', 'Obligatoire par consigne TFE, intégration native Render/Vercel, suivi de projet'),
]
for row in tech_rows:
    add_row(tbl, row)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  8. DIAGRAMME DE CLASSES
# ══════════════════════════════════════════════════════════════════════════════
h1('8. Diagramme de Classes')
body(
    "Le diagramme de classes représente les 15 entités du domaine métier avec leurs "
    "attributs, types, associations et 7 types ENUM. "
    "Insérer le code PlantUML ci-dessous pour générer le diagramme."
)

plantuml_class = '''@startuml Diagramme_Classes_Gestion_Snack
skinparam classBackgroundColor #EEF6FF
skinparam classBorderColor #007A8A
skinparam arrowColor #1A235E
skinparam stereotypeCBackgroundColor #FFD580
skinparam classAttributeIconSize 0

' ── ENUMs ──────────────────────────────────────────────────────────────
enum RoleType {
  ADMIN
  CASHIER
  WAITER
  COOK
  CUSTOMER
  PROVIDER
}

enum ProductType {
  FOOD
  DRINK
  INGREDIENT
}

enum TableStatusType {
  FREE
  OCCUPIED
  RESERVED
}

enum OrderStatusType {
  ACTIVE
  CLOSED
  CANCELLED
}

enum OrderTypeType {
  ON_SITE
  TAKEAWAY
}

enum PaymentMethodType {
  CASH
  CARD
  STRIPE
}

' ── Entités ──────────────────────────────────────────────────────────
class User {
  - id         : Long
  - name       : String
  - email      : String  <<unique>>
  - password   : String  <<BCrypt>>
  - role       : RoleType
  - active     : Boolean
  - createdAt  : LocalDateTime
  --
  + login(email: String, password: String) : Boolean
  + logout() : void
  + updatePassword(newPwd: String) : void
  + deactivate() : void
  + hasRole(role: RoleType) : Boolean
}

class Employee {
  - id        : Long
  - phone     : String
  - position  : String
  - hireDate  : LocalDate
  --
  + getAssignedOrders() : List<Order>
  + processPayment(order: Order, method: PaymentMethodType) : Transaction
}

class Customer {
  - id             : Long
  - loyaltyPoints  : Integer
  --
  + placeOrder(items: List<OrderItem>, tableId: Long) : Order
  + makeReservation(date: LocalDate, time: LocalTime, guests: Integer) : Reservation
  + leaveReview(order: Order, rating: Integer, comment: String) : Review
  + getOrderHistory() : List<Order>
  + addLoyaltyPoints(points: Integer) : void
}

class Provider {
  - id           : Long
  - companyName  : String
  - phone        : String
  - address      : String
  --
  + getSupplyOrders() : List<Order>
  + confirmDelivery(orderId: Long) : void
  + getProductsInAlert() : List<Product>
}

class Product {
  - id             : Long
  - name           : String
  - description    : String
  - price          : BigDecimal
  - type           : ProductType
  - stock          : Integer
  - alertThreshold : Integer
  - imageUrl       : String
  - active         : Boolean
  --
  + decrementStock(qty: Integer) : void
  + incrementStock(qty: Integer) : void
  + isAvailable() : Boolean
  + isStockCritical() : Boolean
  + activate() : void
  + deactivate() : void
}

class Order {
  - id        : Long
  - orderType : OrderTypeType
  - status    : OrderStatusType
  - total     : BigDecimal
  - createdAt : LocalDateTime
  - closedAt  : LocalDateTime
  --
  + addItem(product: Product, qty: Integer) : OrderItem
  + removeItem(itemId: Long) : void
  + calculateTotal() : BigDecimal
  + close() : void
  + cancel() : void
  + isEditable() : Boolean
}

class OrderItem {
  - id        : Long
  - quantity  : Integer
  - unitPrice : BigDecimal
  --
  + getSubtotal() : BigDecimal
}

class TableSnack {
  - id       : Long
  - number   : Integer
  - capacity : Integer
  - status   : TableStatusType
  --
  + occupy() : void
  + free() : void
  + reserve() : void
  + isAvailable() : Boolean
}

class Reservation {
  - id              : Long
  - reservationDate : LocalDate
  - reservationTime : LocalTime
  - numberOfGuests  : Integer
  - customerName    : String
  - customerPhone   : String
  - status          : String
  --
  + confirm() : void
  + cancel() : void
  + isUpcoming() : Boolean
}

class Transaction {
  - id            : Long
  - amount        : BigDecimal
  - paymentMethod : PaymentMethodType
  - stripeId      : String
  - createdAt     : LocalDateTime
  --
  + process() : void
  + refund() : void
  + isStripePayment() : Boolean
}

class Review {
  - id        : Long
  - rating    : Integer
  - comment   : String
  - createdAt : LocalDateTime
  --
  + submit() : void
  + isValid() : Boolean
}

class StockAlert {
  - id        : Long
  - message   : String
  - resolved  : Boolean
  - createdAt : LocalDateTime
  --
  + resolve() : void
  + isActive() : Boolean
}

class StockMovement {
  - id           : Long
  - quantity     : Integer
  - movementType : String
  - reason       : String
  - createdAt    : LocalDateTime
  --
  + record() : void
  + getImpact() : Integer
}

class AuditLog {
  - id         : Long
  - action     : String
  - entityType : String
  - entityId   : Long
  - details    : String
  - timestamp  : LocalDateTime
  --
  + log(action: String, entity: String, user: User) : void
  + toJson() : String
}

class Message {
  - id        : Long
  - content   : String
  - sender    : String
  - createdAt : LocalDateTime
  --
  + send() : void
  + broadcast(topic: String) : void
}

' ── Relations ─────────────────────────────────────────────────────────
User         "1" *-- "0..1" Employee     : est-un >
User         "1" *-- "0..1" Customer     : est-un >
User         "1" *-- "0..1" Provider     : est-un >

Product      "0..*" --> "1"   Provider   : fourni par >
Order        "0..*" --> "0..1" Customer  : passée par >
Order        "0..*" --> "0..1" TableSnack : sur >
Order        "0..*" --> "0..1" Employee  : servie par >
Order        "1"    *-- "1..*" OrderItem  : contient >
OrderItem    "0..*" --> "1"   Product    : concerne >

Transaction  "0..*" --> "1"   Order      : règle >
Transaction  "0..*" --> "0..1" Employee  : traitée par >

Reservation  "0..*" --> "1"   TableSnack : pour >
Reservation  "0..*" --> "0..1" Customer  : faite par >

Review       "0..*" --> "1"   Customer   : écrite par >
Review       "0..*" --> "1"   Order      : concerne >

StockAlert   "0..*" --> "1"   Product    : concerne >
StockMovement "0..*" --> "1"  Product    : concerne >
StockMovement "0..*" --> "0..1" Employee : effectuée par >
AuditLog     "0..*" --> "0..1" User      : imputée à >

' ── ENUMs liés ─────────────────────────────────────────────────────────
User        ..> RoleType
Product     ..> ProductType
TableSnack  ..> TableStatusType
Order       ..> OrderStatusType
Order       ..> OrderTypeType
Transaction ..> PaymentMethodType

@enduml'''

h2('8.1 Code PlantUML — Diagramme de Classes')
code_block(plantuml_class)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  9. MCD / MPD
# ══════════════════════════════════════════════════════════════════════════════
h1('9. Modèle Conceptuel et Physique de Données')

h2('9.1 Synthèse des entités et associations')
body(
    "Le modèle de données repose sur 14 tables en base de données. "
    "Les héritages sont implémentés via des relations 1-1 (users ↔ employees/customers/providers). "
    "Les ENUMs PostgreSQL natifs assurent l\'intégrité des valeurs discrètes."
)

h2('9.2 Tables et cardinalités principales')
tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
table_header(tbl, ['Entité', 'Clé primaire', 'Relations principales'])
mcd_rows = [
    ('users',        'id (BIGSERIAL)', '1–0..1 employees, 1–0..1 customers, 1–0..1 providers'),
    ('employees',    'id (BIGSERIAL)', 'FK user_id → users, référencé par orders, transactions, stock_movements'),
    ('customers',    'id (BIGSERIAL)', 'FK user_id → users, référencé par orders, reservations, reviews'),
    ('providers',    'id (BIGSERIAL)', 'FK user_id → users, référencé par products'),
    ('products',     'id (BIGSERIAL)', 'FK provider_id → providers, référencé par order_items, stock_alerts, stock_movements'),
    ('orders',       'id (BIGSERIAL)', 'FK customer_id, table_id, employee_id ; 1-N order_items, 1-1 transactions, 1-N reviews'),
    ('order_items',  'id (BIGSERIAL)', 'FK order_id → orders, FK product_id → products'),
    ('tables_snack', 'id (BIGSERIAL)', 'Référencé par orders, reservations'),
    ('reservations', 'id (BIGSERIAL)', 'FK table_id → tables_snack, FK customer_id → customers'),
    ('transactions', 'id (BIGSERIAL)', 'FK order_id → orders, FK employee_id → employees'),
    ('reviews',      'id (BIGSERIAL)', 'FK order_id → orders, FK customer_id → customers'),
    ('stock_alerts', 'id (BIGSERIAL)', 'FK product_id → products'),
    ('stock_movements','id (BIGSERIAL)','FK product_id → products, FK employee_id → employees'),
    ('audit_log',    'id (BIGSERIAL)', 'FK user_id → users'),
]
for row in mcd_rows:
    add_row(tbl, row)

doc.add_paragraph()

# ─── MCD PlantUML ────────────────────────────────────────────────────────────
plantuml_mcd = '''@startuml MCD_Gestion_Snack
skinparam entityBackgroundColor #EEF6FF
skinparam entityBorderColor #007A8A
skinparam arrowColor #1A235E
skinparam linetype ortho
hide empty methods

' ── Entités ────────────────────────────────────────────────────
entity "UTILISATEUR" as USR {
  nom
  email
  mot_de_passe
  role
  actif
  date_creation
}

entity "EMPLOYE" as EMP {
  telephone
  poste
  date_embauche
}

entity "CLIENT" as CLT {
  points_fidelite
}

entity "FOURNISSEUR" as PRV {
  nom_societe
  telephone
  adresse
}

entity "PRODUIT" as PRD {
  nom
  description
  prix
  type
  stock
  seuil_alerte
  actif
}

entity "COMMANDE" as CMD {
  type_commande
  statut
  total
  date_creation
  date_cloture
}

entity "ARTICLE_COMMANDE" as ART {
  quantite
  prix_unitaire
}

entity "TABLE" as TBL {
  numero
  capacite
  statut
}

entity "RESERVATION" as RES {
  date_reservation
  heure_reservation
  nb_couverts
  nom_client
  telephone_client
  statut
}

entity "TRANSACTION" as TRX {
  montant
  methode_paiement
  id_stripe
  date
}

entity "AVIS" as AVI {
  note
  commentaire
  date
}

entity "ALERTE_STOCK" as ALR {
  message
  resolu
  date
}

entity "MOUVEMENT_STOCK" as MVT {
  quantite
  type_mouvement
  motif
  date
}

entity "JOURNAL_AUDIT" as AUD {
  action
  type_entite
  id_entite
  details
  horodatage
}

' ── Generalisation / Specialisation ────────────────────────────
' (1,1) cote UTILISATEUR — (0,1) cote specialisation
USR ||--o| EMP : "est un"
USR ||--o| CLT : "est un"
USR ||--o| PRV : "est un"

' ── Associations ───────────────────────────────────────────────
PRV ||--o{ PRD   : "fournit (0,n)"
CMD ||--|{ ART   : "contient (1,n)"
PRD ||--o{ ART   : "figure dans (0,n)"
CLT |o--o{ CMD   : "passe (0,n)"
EMP |o--o{ CMD   : "gere (0,n)"
TBL |o--o{ CMD   : "accueille (0,n)"
TBL ||--o{ RES   : "fait l objet de (0,n)"
CLT |o--o{ RES   : "effectue (0,n)"
CMD ||--o| TRX   : "est reglee par (0,1)"
EMP |o--o{ TRX   : "traite (0,n)"
CLT |o--o{ AVI   : "redige (0,n)"
CMD ||--o| AVI   : "fait l objet de (0,1)"
PRD ||--o{ ALR   : "genere (0,n)"
PRD ||--o{ MVT   : "concerne (0,n)"
EMP |o--o{ MVT   : "effectue (0,n)"
USR |o--o{ AUD   : "genere (0,n)"

@enduml'''

h2('9.3 Code PlantUML — Modèle Conceptuel de Données (MCD)')
body(
    "Le MCD représente les entités du domaine métier avec leurs attributs conceptuels "
    "et leurs associations avec les cardinalités Merise. Les clés primaires et étrangères "
    "n'apparaissent pas ici — elles relèvent du MPD."
)
code_block(plantuml_mcd)

page_break()

# ─── MPD PlantUML — Partie 1 ─────────────────────────────────────────────────
plantuml_mpd1 = '''@startuml MPD_Partie1_Identites_Catalogue
skinparam entityBackgroundColor #EEF6FF
skinparam entityBorderColor #007A8A
skinparam arrowColor #1A235E
skinparam linetype ortho
hide empty methods

' Legende : * = PK  |  # = FK  |  NN = NOT NULL  |  UQ = UNIQUE

entity "users" as USR {
  * id : BIGSERIAL
  --
  name       : VARCHAR(100) NN
  email      : VARCHAR(150) NN UQ
  password   : VARCHAR(255) NN
  role       : role_type NN
  active     : BOOLEAN
  created_at : TIMESTAMP
}

entity "employees" as EMP {
  * id : BIGSERIAL
  --
  # user_id  : BIGINT UQ NN
  phone      : VARCHAR(20)
  position   : VARCHAR(100)
  hire_date  : DATE
}

entity "customers" as CLT {
  * id : BIGSERIAL
  --
  # user_id       : BIGINT UQ NN
  loyalty_points  : INTEGER
}

entity "providers" as PRV {
  * id : BIGSERIAL
  --
  # user_id     : BIGINT UQ NN
  company_name  : VARCHAR(200) NN
  phone         : VARCHAR(20)
  address       : TEXT
}

entity "products" as PRD {
  * id : BIGSERIAL
  --
  # provider_id   : BIGINT
  name            : VARCHAR(200) NN
  description     : TEXT
  price           : NUMERIC(10,2) NN
  type            : product_type NN
  stock           : INTEGER
  alert_threshold : INTEGER
  image_url       : TEXT
  active          : BOOLEAN
}

entity "orders" as CMD {
  * id : BIGSERIAL
  --
  # customer_id : BIGINT
  # table_id    : BIGINT
  # employee_id : BIGINT
  order_type    : order_type_type NN
  status        : order_status_type
  total         : NUMERIC(10,2)
  created_at    : TIMESTAMP
  closed_at     : TIMESTAMP
}

entity "order_items" as ART {
  * id : BIGSERIAL
  --
  # order_id   : BIGINT NN
  # product_id : BIGINT NN
  quantity     : INTEGER NN
  unit_price   : NUMERIC(10,2) NN
}

' ── FK ─────────────────────────────────────────────────────────
USR ||--o| EMP : "user_id"
USR ||--o| CLT : "user_id"
USR ||--o| PRV : "user_id"
PRV ||--o{ PRD : "provider_id"
CLT |o--o{ CMD : "customer_id"
EMP |o--o{ CMD : "employee_id"
CMD ||--|{ ART : "order_id"
PRD ||--o{ ART : "product_id"

@enduml'''

# ─── MPD PlantUML — Partie 2 ─────────────────────────────────────────────────
plantuml_mpd2 = '''@startuml MPD_Partie2_Operations_Support
skinparam entityBackgroundColor #EEF6FF
skinparam entityBorderColor #007A8A
skinparam arrowColor #1A235E
skinparam linetype ortho
hide empty methods

' Legende : * = PK  |  # = FK  |  NN = NOT NULL  |  UQ = UNIQUE
' Entites en gris clair = references vers la Partie 1

skinparam entity {
  BackgroundColor<<ref>> #F5F5F5
  BorderColor<<ref>> #BBBBBB
}

entity "users <<ref>>" as USR <<ref>> {
  * id : BIGSERIAL
}
entity "orders <<ref>>" as CMD <<ref>> {
  * id : BIGSERIAL
}
entity "customers <<ref>>" as CLT <<ref>> {
  * id : BIGSERIAL
}
entity "employees <<ref>>" as EMP <<ref>> {
  * id : BIGSERIAL
}
entity "products <<ref>>" as PRD <<ref>> {
  * id : BIGSERIAL
}

entity "tables_snack" as TBL {
  * id : BIGSERIAL
  --
  number   : INTEGER UQ NN
  capacity : INTEGER NN
  status   : table_status_type
}

entity "reservations" as RES {
  * id : BIGSERIAL
  --
  # table_id       : BIGINT NN
  # customer_id    : BIGINT
  reservation_date : DATE NN
  reservation_time : TIME NN
  number_of_guests : INTEGER NN
  customer_name    : VARCHAR(100) NN
  customer_phone   : VARCHAR(20)
  status           : VARCHAR(50)
}

entity "transactions" as TRX {
  * id : BIGSERIAL
  --
  # order_id      : BIGINT NN
  # employee_id   : BIGINT
  amount          : NUMERIC(10,2) NN
  payment_method  : payment_method_type NN
  stripe_id       : VARCHAR(255)
  created_at      : TIMESTAMP
}

entity "reviews" as AVI {
  * id : BIGSERIAL
  --
  # customer_id : BIGINT NN
  # order_id    : BIGINT NN
  rating        : INTEGER NN
  comment       : TEXT
  created_at    : TIMESTAMP
}

entity "stock_alerts" as ALR {
  * id : BIGSERIAL
  --
  # product_id : BIGINT NN
  message      : TEXT NN
  resolved     : BOOLEAN
  created_at   : TIMESTAMP
}

entity "stock_movements" as MVT {
  * id : BIGSERIAL
  --
  # product_id  : BIGINT NN
  # employee_id : BIGINT
  quantity      : INTEGER NN
  movement_type : VARCHAR(50) NN
  reason        : TEXT
  created_at    : TIMESTAMP
}

entity "audit_log" as AUD {
  * id : BIGSERIAL
  --
  # user_id    : BIGINT
  action       : VARCHAR(100) NN
  entity_type  : VARCHAR(100) NN
  entity_id    : BIGINT
  details      : TEXT
  timestamp    : TIMESTAMP
}

' ── FK ─────────────────────────────────────────────────────────
TBL ||--o{ RES  : "table_id"
CLT |o--o{ RES  : "customer_id"
CMD ||--o| TRX  : "order_id"
EMP |o--o{ TRX  : "employee_id"
CLT ||--o{ AVI  : "customer_id"
CMD ||--o| AVI  : "order_id"
PRD ||--o{ ALR  : "product_id"
PRD ||--o{ MVT  : "product_id"
EMP |o--o{ MVT  : "employee_id"
USR |o--o{ AUD  : "user_id"

@enduml'''

h2('9.4 Code PlantUML — MPD Partie 1 : Identités & Catalogue')
body(
    "Tables : users, employees, customers, providers, products, orders, order_items. "
    "Légende : * = PK, # = FK, NN = NOT NULL, UQ = UNIQUE."
)
code_block(plantuml_mpd1)

doc.add_paragraph()
h2('9.5 Code PlantUML — MPD Partie 2 : Opérations & Support')
body(
    "Tables : tables_snack, reservations, transactions, reviews, stock_alerts, stock_movements, audit_log. "
    "Les entités grisées (<<ref>>) sont des références vers les tables de la Partie 1."
)
code_block(plantuml_mpd2)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  10. DICTIONNAIRE DE DONNÉES
# ══════════════════════════════════════════════════════════════════════════════
h1('10. Dictionnaire de Données')
body("Dictionnaire complet des 14 tables de la base de données PostgreSQL.")

def dict_table(name: str, desc: str, columns: list):
    h2(f'Table : {name}')
    body(desc)
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = 'Table Grid'
    table_header(tbl, ['Attribut', 'Type SQL', 'Contrainte', 'Description', 'Exemple'])
    for col in columns:
        add_row(tbl, col)
    doc.add_paragraph()

dict_table('users', 'Table centrale des utilisateurs du système (tous rôles confondus).', [
    ('id',         'BIGSERIAL',     'PK NOT NULL',      'Identifiant unique auto-incrémenté',     '1'),
    ('name',       'VARCHAR(100)',  'NOT NULL',         'Nom complet de l\'utilisateur',          'Jean Dupont'),
    ('email',      'VARCHAR(150)',  'UNIQUE NOT NULL',  'Adresse email (identifiant de connexion)','jean@snack.be'),
    ('password',   'VARCHAR(255)',  'NOT NULL',         'Mot de passe haché BCrypt',              '$2a$10$...'),
    ('role',       'role_type',     'NOT NULL',         'Rôle ENUM : ADMIN, CASHIER, WAITER, COOK, CUSTOMER, PROVIDER', 'ADMIN'),
    ('active',     'BOOLEAN',       'DEFAULT TRUE',     'Compte actif/désactivé',                'true'),
    ('created_at', 'TIMESTAMP',     'DEFAULT NOW()',    'Date et heure de création',              '2025-01-15 10:30:00'),
])

dict_table('employees', 'Informations complémentaires pour les employés (ADMIN, CASHIER, WAITER, COOK).', [
    ('id',       'BIGSERIAL',    'PK NOT NULL',    'Identifiant unique',                  '1'),
    ('user_id',  'BIGINT',       'FK UNIQUE users','Référence vers users.id',             '1'),
    ('phone',    'VARCHAR(20)',  'NULL',           'Numéro de téléphone',                 '+32 470 123 456'),
    ('position', 'VARCHAR(100)', 'NULL',           'Intitulé du poste',                  'Chef cuisinier'),
    ('hire_date','DATE',         'NULL',           'Date d\'embauche',                   '2024-03-01'),
])

dict_table('customers', 'Informations complémentaires pour les clients.', [
    ('id',             'BIGSERIAL', 'PK NOT NULL',    'Identifiant unique',           '1'),
    ('user_id',        'BIGINT',    'FK UNIQUE users','Référence vers users.id',      '3'),
    ('loyalty_points', 'INTEGER',   'DEFAULT 0',      'Points de fidélité cumulés',   '150'),
])

dict_table('providers', 'Informations complémentaires pour les fournisseurs.', [
    ('id',           'BIGSERIAL',    'PK NOT NULL',    'Identifiant unique',     '1'),
    ('user_id',      'BIGINT',       'FK UNIQUE users','Référence vers users.id','5'),
    ('company_name', 'VARCHAR(200)', 'NOT NULL',       'Nom de la société',      'Metro Belgique SA'),
    ('phone',        'VARCHAR(20)',  'NULL',           'Téléphone professionnel','02 567 89 00'),
    ('address',      'TEXT',         'NULL',           'Adresse de livraison',   'Rue du Commerce 12, Bruxelles'),
])

dict_table('products', 'Catalogue des produits disponibles à la vente ou en ingrédients.', [
    ('id',              'BIGSERIAL',      'PK NOT NULL',    'Identifiant unique',                    '1'),
    ('name',            'VARCHAR(200)',   'NOT NULL',       'Nom du produit',                        'Burger Classic'),
    ('description',     'TEXT',           'NULL',           'Description détaillée',                 'Steak haché, salade, tomate'),
    ('price',           'NUMERIC(10,2)',  'NOT NULL',       'Prix de vente en euros',                '8.50'),
    ('type',            'product_type',   'NOT NULL',       'Type ENUM : FOOD, DRINK, INGREDIENT',   'FOOD'),
    ('stock',           'INTEGER',        'DEFAULT 0',      'Quantité en stock',                     '25'),
    ('alert_threshold', 'INTEGER',        'DEFAULT 5',      'Seuil déclenchant l\'alerte de stock',  '5'),
    ('image_url',       'TEXT',           'NULL',           'URL de l\'image du produit',            'https://...'),
    ('active',          'BOOLEAN',        'DEFAULT TRUE',   'Produit disponible à la vente',         'true'),
    ('provider_id',     'BIGINT',         'FK providers',   'Fournisseur du produit',                '1'),
])

dict_table('orders', 'Table des commandes passées dans le snack.', [
    ('id',          'BIGSERIAL',         'PK NOT NULL',    'Identifiant unique',                                  '1'),
    ('order_type',  'order_type_type',   'NOT NULL',       'Type ENUM : ON_SITE ou TAKEAWAY',                    'ON_SITE'),
    ('status',      'order_status_type', 'DEFAULT ACTIVE', 'Statut ENUM : ACTIVE, CLOSED, CANCELLED',           'ACTIVE'),
    ('total',       'NUMERIC(10,2)',      'DEFAULT 0',      'Montant total de la commande',                       '15.50'),
    ('created_at',  'TIMESTAMP',          'DEFAULT NOW()',  'Date et heure de création',                          '2025-06-10 12:30:00'),
    ('closed_at',   'TIMESTAMP',          'NULL',           'Date et heure de clôture',                           '2025-06-10 13:00:00'),
    ('customer_id', 'BIGINT',             'FK customers',   'Client ayant passé la commande (nullable si espèces)','2'),
    ('table_id',    'BIGINT',             'FK tables_snack','Table concernée (nullable si emporter)',              '3'),
    ('employee_id', 'BIGINT',             'FK employees',   'Employé ayant créé la commande',                     '4'),
])

dict_table('order_items', 'Lignes détaillées de chaque commande.', [
    ('id',         'BIGSERIAL',     'PK NOT NULL',  'Identifiant unique',                '1'),
    ('order_id',   'BIGINT',        'FK orders',    'Référence vers orders.id',          '1'),
    ('product_id', 'BIGINT',        'FK products',  'Référence vers products.id',        '5'),
    ('quantity',   'INTEGER',       'NOT NULL',     'Quantité commandée',                '2'),
    ('unit_price', 'NUMERIC(10,2)', 'NOT NULL',     'Prix unitaire au moment de la commande','8.50'),
])

dict_table('tables_snack', 'Tables physiques du restaurant.', [
    ('id',       'BIGSERIAL',          'PK NOT NULL',     'Identifiant unique',                        '1'),
    ('number',   'INTEGER',            'UNIQUE NOT NULL', 'Numéro de la table',                        '5'),
    ('capacity', 'INTEGER',            'NOT NULL',        'Nombre de places assises',                  '4'),
    ('status',   'table_status_type',  'DEFAULT FREE',    'Statut ENUM : FREE, OCCUPIED, RESERVED',    'FREE'),
])

dict_table('reservations', 'Réservations de tables effectuées par les clients.', [
    ('id',               'BIGSERIAL',   'PK NOT NULL',   'Identifiant unique',                '1'),
    ('reservation_date', 'DATE',         'NOT NULL',      'Date de la réservation',            '2025-07-01'),
    ('reservation_time', 'TIME',         'NOT NULL',      'Heure de la réservation',           '19:30'),
    ('number_of_guests', 'INTEGER',      'NOT NULL',      'Nombre de convives',                '4'),
    ('customer_name',    'VARCHAR(100)', 'NOT NULL',      'Nom du client (peut être anonyme)', 'Marie Lambert'),
    ('customer_phone',   'VARCHAR(20)',  'NULL',          'Téléphone du client',               '+32 478 000 000'),
    ('status',           'VARCHAR(50)',  'DEFAULT \'PENDING\'', 'Statut : PENDING, CONFIRMED, CANCELLED','CONFIRMED'),
    ('table_id',         'BIGINT',       'FK tables_snack','Table réservée',                   '3'),
    ('customer_id',      'BIGINT',       'FK customers',  'Client ayant réservé (nullable)',    '2'),
])

dict_table('transactions', 'Enregistrement de chaque paiement effectué.', [
    ('id',             'BIGSERIAL',            'PK NOT NULL',  'Identifiant unique',                        '1'),
    ('amount',         'NUMERIC(10,2)',         'NOT NULL',     'Montant encaissé',                          '15.50'),
    ('payment_method', 'payment_method_type',   'NOT NULL',     'Méthode ENUM : CASH, CARD, STRIPE',        'CARD'),
    ('stripe_id',      'VARCHAR(255)',           'NULL',         'Identifiant de paiement Stripe (si STRIPE)','pi_3Xxx...'),
    ('created_at',     'TIMESTAMP',              'DEFAULT NOW()','Horodatage du paiement',                  '2025-06-10 13:00:00'),
    ('order_id',       'BIGINT',                 'FK orders',    'Commande réglée',                         '1'),
    ('employee_id',    'BIGINT',                 'FK employees', 'Caissier ayant encaissé',                 '2'),
])

dict_table('reviews', 'Avis et évaluations laissés par les clients.', [
    ('id',          'BIGSERIAL',  'PK NOT NULL',  'Identifiant unique',               '1'),
    ('rating',      'INTEGER',    'NOT NULL CHECK (rating BETWEEN 1 AND 5)', 'Note de 1 à 5 étoiles','5'),
    ('comment',     'TEXT',       'NULL',         'Commentaire libre',                'Excellent burger !'),
    ('created_at',  'TIMESTAMP',  'DEFAULT NOW()','Date et heure de l\'avis',        '2025-06-10 14:00:00'),
    ('customer_id', 'BIGINT',     'FK customers', 'Client auteur de l\'avis',        '2'),
    ('order_id',    'BIGINT',     'FK orders',    'Commande évaluée',                '1'),
])

dict_table('stock_alerts', 'Alertes générées automatiquement par trigger PostgreSQL.', [
    ('id',         'BIGSERIAL',  'PK NOT NULL',   'Identifiant unique',                         '1'),
    ('message',    'TEXT',       'NOT NULL',      'Description de l\'alerte',                   'Stock Coca Cola : 3 (seuil : 5)'),
    ('resolved',   'BOOLEAN',    'DEFAULT FALSE', 'Alerte traitée ou non',                      'false'),
    ('created_at', 'TIMESTAMP',  'DEFAULT NOW()','Date et heure de génération de l\'alerte',   '2025-06-10 12:35:00'),
    ('product_id', 'BIGINT',     'FK products',  'Produit en rupture ou sous le seuil',        '8'),
])

dict_table('audit_log', 'Journal de toutes les actions utilisateurs pour traçabilité.', [
    ('id',          'BIGSERIAL',  'PK NOT NULL',  'Identifiant unique',                      '1'),
    ('action',      'VARCHAR(100)','NOT NULL',    'Type d\'action : CREATE, UPDATE, DELETE, LOGIN, etc.','UPDATE'),
    ('entity_type', 'VARCHAR(100)','NOT NULL',    'Entité concernée',                        'Product'),
    ('entity_id',   'BIGINT',     'NULL',         'Identifiant de l\'entité concernée',      '5'),
    ('details',     'TEXT',       'NULL',         'Détails supplémentaires (JSON ou texte)', '{"field":"price","old":"8.00","new":"9.00"}'),
    ('timestamp',   'TIMESTAMP',  'DEFAULT NOW()','Horodatage de l\'action',                '2025-06-10 11:00:00'),
    ('user_id',     'BIGINT',     'FK users',     'Utilisateur responsable de l\'action',   '1'),
])

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  11. TRIGGERS POSTGRESQL
# ══════════════════════════════════════════════════════════════════════════════
h1('11. Triggers PostgreSQL')
body(
    "Deux triggers automatiques garantissent la cohérence des stocks sans intervention humaine."
)

h2('11.1 trigger_update_stock')
body("Déclencheur : AFTER INSERT sur order_items.")
body("Effet : décrémente automatiquement products.stock de la quantité commandée.")
bullet("Garantit que chaque article commandé est immédiatement déduit du stock.")
bullet("Évite les incohérences dues à des mises à jour manuelles oubliées.")

h2('11.2 trigger_stock_alert')
body("Déclencheur : AFTER UPDATE sur products (colonne stock).")
body("Effet : insère une ligne dans stock_alerts si stock < alert_threshold.")
bullet("Génère automatiquement une alerte dès que le stock devient critique.")
bullet("Le backend WebSocket détecte la nouvelle alerte et la diffuse en temps réel.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  12. DIAGRAMMES DE SÉQUENCE
# ══════════════════════════════════════════════════════════════════════════════
h1('12. Diagrammes de Séquence')
body(
    "Cinq scénarios critiques sont modélisés. Les rectangles d\'activation verticaux "
    "(générés par les commandes activate/deactivate PlantUML) représentent uniquement "
    "la période d\'activité de chaque participant, pas toute la ligne de vie."
)

# ── Séquence 1 — Connexion ────────────────────────────────────────────────────
h2('12.1 Séquence 1 — Authentification / Connexion')
body(
    "Scénario : un utilisateur saisit ses identifiants, le backend vérifie le mot de passe "
    "haché BCrypt et retourne les informations de session."
)

plantuml_seq1 = '''@startuml Sequence_Connexion
skinparam sequenceArrowThickness 1.5
skinparam sequenceParticipantBorderColor #007A8A
skinparam sequenceLifeLineBorderColor #AAAAAA
skinparam sequenceBoxBorderColor #1A235E
skinparam noteBorderColor #E67E22

actor "Utilisateur" as U
participant "Interface Web\\n(React)" as IW
participant "AuthController\\n(Spring Boot)" as AC
database "PostgreSQL\\n(Neon.tech)" as BD

U -> IW : Saisir email + mot de passe
activate IW

IW -> AC : POST /api/auth/login\\n{email, password}
activate AC

AC -> BD : SELECT * FROM users WHERE email = ?
activate BD
BD --> AC : Retourne l'utilisateur (hash BCrypt)
deactivate BD

AC -> AC : BCrypt.matches(password, hash)
activate AC
AC --> AC
deactivate AC

alt Succès — mot de passe correct
  AC --> IW : 200 OK — {userId, name, role}
  deactivate AC
  IW -> IW : Stocker session (localStorage/context)
  activate IW
  IW --> IW
  deactivate IW
  IW --> U : Rediriger vers tableau de bord
else Échec — mot de passe incorrect
  AC --> IW : 401 Unauthorized
  deactivate AC
  IW --> U : Afficher "Email ou mot de passe incorrect"
end

deactivate IW
@enduml'''

h3('Code PlantUML — Séquence Connexion')
code_block(plantuml_seq1)

# ── Séquence 2 — Prise de commande ────────────────────────────────────────────
h2('12.2 Séquence 2 — Prise de Commande (Serveur)')
body(
    "Scénario : le serveur crée une commande pour une table, le trigger décrémente le stock "
    "et une notification WebSocket est envoyée au cuisinier."
)

plantuml_seq2 = '''@startuml Sequence_Commande
skinparam sequenceArrowThickness 1.5
skinparam sequenceParticipantBorderColor #007A8A
skinparam sequenceLifeLineBorderColor #AAAAAA

actor "Serveur" as W
participant "Interface Web\\n(React)" as IW
participant "OrderController\\n(Spring Boot)" as OC
database "PostgreSQL\\n(Neon.tech)" as BD
participant "WebSocket\\nBroker" as WS
actor "Cuisinier" as K

W -> IW : Sélectionner table + articles
activate IW

IW -> OC : POST /api/orders\\n{tableId, items[], employeeId}
activate OC

OC -> BD : INSERT INTO orders (...)
activate BD
BD --> OC : orderId = 42
deactivate BD

OC -> BD : INSERT INTO order_items (orderId, productId, qty, ...)
activate BD
BD -> BD : trigger_update_stock\\n→ UPDATE products SET stock = stock - qty
activate BD
BD --> BD
deactivate BD
BD --> OC : order_items créés
deactivate BD

OC -> WS : /topic/orders — nouvelle commande #42
activate WS
WS --> K : Notification "Nouvelle commande #42"
deactivate WS

OC --> IW : 201 Created — {orderId: 42, status: ACTIVE}
deactivate OC

IW --> W : Afficher confirmation commande
deactivate IW

K -> IW : Consulter détail commande #42
activate IW
IW -> OC : GET /api/orders/42
activate OC
OC -> BD : SELECT order + items WHERE id = 42
activate BD
BD --> OC : Commande complète
deactivate BD
OC --> IW : 200 OK — détail commande
deactivate OC
IW --> K : Afficher articles à préparer
deactivate IW

@enduml'''

h3('Code PlantUML — Séquence Prise de Commande')
code_block(plantuml_seq2)

# ── Séquence 3 — Chatbot Vocal ───────────────────────────────────────────────
h2('12.3 Séquence 3 — Chatbot Vocal (Client)')
body(
    "Scénario : le client active le microphone, parle au chatbot. Le système reconnaît la voix, "
    "génère une réponse IA via Groq/LLaMA 3.1, puis lit la réponse avec ElevenLabs TTS."
)

plantuml_seq3 = '''@startuml Sequence_Chatbot_Vocal
skinparam sequenceArrowThickness 1.5
skinparam sequenceParticipantBorderColor #007A8A
skinparam sequenceLifeLineBorderColor #AAAAAA

actor "Client" as C
participant "Interface Web\\n(React)" as IW
participant "Web Speech API\\n(Navigateur)" as STT
participant "Groq API\\n(LLaMA 3.1-8b)" as GROQ
participant "ElevenLabs API\\n(TTS)" as EL
participant "AudioContext\\n(Navigateur)" as AUD

C -> IW : Cliquer bouton microphone
activate IW

IW -> STT : SpeechRecognition.start()\\nlang: fr-FR, continuous: false
activate STT
STT --> C : Écoute active (indicateur visuel)

C -> STT : Parler ("Qu'est-ce qu'il y a au menu ?")
STT --> IW : onresult — transcript: "Qu'est-ce qu'il y a au menu ?"
deactivate STT

IW -> IW : État: LISTENING → PROCESSING
activate IW
IW --> IW
deactivate IW

IW -> GROQ : POST /openai/v1/chat/completions\\n{model: llama-3.1-8b-instant,\\n messages: [system_prompt, user_msg]}
activate GROQ
GROQ --> IW : 200 OK — {choices[0].message.content: "Voici notre menu..."}
deactivate GROQ

IW -> IW : État: PROCESSING → SPEAKING
activate IW
IW --> IW
deactivate IW

IW -> EL : POST /v1/text-to-speech/pNInz6obpgDQGcFmaJgB\\n{text: réponse, model_id: eleven_turbo_v2_5}
activate EL
EL --> IW : 200 OK — audio/mpeg (flux audio)
deactivate EL

IW -> AUD : new Audio(blobURL).play()
activate AUD
AUD --> C : Lecture vocale de la réponse
deactivate AUD

IW -> IW : État: SPEAKING → LISTENING (après fin audio)
activate IW
IW --> IW
deactivate IW

IW --> C : Afficher réponse textuelle + prêt pour nouvelle question
deactivate IW

note over IW : Fallback : si ElevenLabs indisponible\\n→ speechSynthesis.speak() (navigateur)

@enduml'''

h3('Code PlantUML — Séquence Chatbot Vocal')
code_block(plantuml_seq3)

# ── Séquence 4 — Paiement ────────────────────────────────────────────────────
h2('12.4 Séquence 4 — Paiement et Génération de Ticket (Caissier)')
body(
    "Scénario : le caissier encaisse une commande par carte Stripe, "
    "la commande est clôturée et un ticket PDF est généré."
)

plantuml_seq4 = '''@startuml Sequence_Paiement
skinparam sequenceArrowThickness 1.5
skinparam sequenceParticipantBorderColor #007A8A
skinparam sequenceLifeLineBorderColor #AAAAAA

actor "Caissier" as CSH
participant "Interface Web\\n(React)" as IW
participant "PaymentController\\n(Spring Boot)" as PC
participant "Stripe API" as STR
database "PostgreSQL\\n(Neon.tech)" as BD
participant "jsPDF\\n(React)" as PDF

CSH -> IW : Sélectionner commande + méthode STRIPE
activate IW

IW -> PC : POST /api/payments/stripe\\n{orderId, amount}
activate PC

PC -> STR : Créer PaymentIntent\\n{amount, currency: eur}
activate STR
STR --> PC : {client_secret, paymentIntentId}
deactivate STR

PC --> IW : 200 OK — {clientSecret}
deactivate PC

IW -> STR : stripe.confirmCardPayment(clientSecret, cardElement)
activate STR
STR --> IW : {status: "succeeded", paymentIntentId}
deactivate STR

IW -> PC : POST /api/payments/confirm\\n{orderId, paymentIntentId}
activate PC

PC -> BD : UPDATE orders SET status = CLOSED WHERE id = ?
activate BD
BD --> PC : OK
deactivate BD

PC -> BD : INSERT INTO transactions\\n{orderId, amount, STRIPE, paymentIntentId}
activate BD
BD --> PC : transactionId
deactivate BD

PC --> IW : 200 OK — {transactionId, status: CLOSED}
deactivate PC

IW -> PDF : new jsPDF()\\nautoTable({order details})
activate PDF
PDF --> IW : Blob PDF généré
deactivate PDF

IW --> CSH : Télécharger / imprimer ticket PDF
deactivate IW

@enduml'''

h3('Code PlantUML — Séquence Paiement')
code_block(plantuml_seq4)

# ── Séquence 5 — Stock automatique ───────────────────────────────────────────
h2('12.5 Séquence 5 — Gestion Automatique du Stock (Trigger)')
body(
    "Scénario : l'insertion d'un article en commande déclenche le trigger de stock, "
    "qui génère une alerte transmise en temps réel à l'administrateur."
)

plantuml_seq5 = '''@startuml Sequence_Stock_Auto
skinparam sequenceArrowThickness 1.5
skinparam sequenceParticipantBorderColor #007A8A
skinparam sequenceLifeLineBorderColor #AAAAAA

participant "OrderService\\n(Spring Boot)" as OS
database "PostgreSQL\\n(Neon.tech)" as BD
participant "trigger_update_stock\\n(PostgreSQL)" as T1
participant "trigger_stock_alert\\n(PostgreSQL)" as T2
participant "StockAlertJob\\n(Spring Boot)" as JOB
participant "WebSocket Broker\\n(Spring Boot)" as WS
actor "Administrateur" as ADM

OS -> BD : INSERT INTO order_items\\n(orderId=42, productId=8, qty=3)
activate BD

BD -> T1 : AFTER INSERT trigger
activate T1
T1 -> BD : UPDATE products\\nSET stock = stock - 3\\nWHERE id = 8
activate BD
BD --> T1 : stock = 2 (était 5, seuil = 5)
deactivate BD
deactivate T1

BD -> T2 : AFTER UPDATE trigger\\n(stock 5 → 2 < alert_threshold=5)
activate T2
T2 -> BD : INSERT INTO stock_alerts\\n(productId=8, message="Stock Coca: 2 < seuil 5")
activate BD
BD --> T2 : alertId = 77
deactivate BD
deactivate T2

BD --> OS : INSERT order_items OK
deactivate BD

JOB -> BD : SELECT * FROM stock_alerts WHERE resolved = false
activate BD
BD --> JOB : [alerte #77 — Coca Cola]
deactivate BD

JOB -> WS : convertAndSend("/topic/alerts", alerte #77)
activate WS
WS --> ADM : Notification WebSocket — "Stock Coca Cola critique : 2 unités"
deactivate WS

ADM -> WS : Marquer comme traitée (resolved = true)
activate WS
WS -> BD : UPDATE stock_alerts SET resolved = true WHERE id = 77
activate BD
BD --> WS : OK
deactivate BD
WS --> ADM : Alerte résolue
deactivate WS

@enduml'''

h3('Code PlantUML — Séquence Stock Automatique')
code_block(plantuml_seq5)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  13. SÉCURITÉ
# ══════════════════════════════════════════════════════════════════════════════
h1('13. Analyse de la Sécurité')

h2('13.1 Authentification et sessions')
bullet("Les mots de passe sont hachés avec BCrypt (facteur de coût 10) via Spring Security Crypto.")
bullet("Aucun mot de passe n'est stocké en clair en base de données.")
bullet("Les sessions sont gérées côté client (stockage des informations utilisateur dans le contexte React).")
bullet("Les endpoints API vérifient le rôle de l'utilisateur avant toute opération sensible.")

h2('13.2 Contrôle d\'accès par rôle (RBAC)')
body("Chaque endpoint Spring Boot est protégé selon le rôle requis :")
tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
table_header(tbl, ['Endpoint', 'Rôle(s) autorisé(s)', 'Description'])
rbac_rows = [
    ('GET /api/products', 'ALL', 'Consulter le catalogue (public)'),
    ('POST /api/products', 'ADMIN', 'Créer un produit'),
    ('PUT /api/products/{id}', 'ADMIN', 'Modifier un produit'),
    ('POST /api/orders', 'WAITER, CUSTOMER', 'Créer une commande'),
    ('GET /api/orders', 'ADMIN, CASHIER, WAITER, COOK', 'Lister les commandes'),
    ('POST /api/payments', 'CASHIER, CUSTOMER', 'Effectuer un paiement'),
    ('GET /api/admin/dashboard', 'ADMIN', 'Tableau de bord analytique'),
    ('GET /api/audit-log', 'ADMIN', 'Journal d\'audit'),
    ('GET /api/users', 'ADMIN', 'Gérer les utilisateurs'),
]
for row in rbac_rows:
    add_row(tbl, row)

doc.add_paragraph()
h2('13.3 Sécurité des communications')
bullet("Toutes les communications entre le client et le serveur transitent en HTTPS (TLS 1.2+).")
bullet("Les certificats SSL sont gérés automatiquement par Render et Vercel.")
bullet("Les clés API (Groq, ElevenLabs, Stripe) sont stockées dans des variables d\'environnement, jamais dans le code source.")
bullet("Le fichier .env est exclu du dépôt Git via .gitignore.")

h2('13.4 Protection contre les injections')
bullet("Spring Data JPA utilise des requêtes paramétrées (PreparedStatement) pour toutes les opérations en base de données.")
bullet("Aucune concaténation de chaînes SQL dans le code.")
bullet("Validation des données entrantes via les annotations Jakarta Validation (@NotNull, @Size, @Email).")

h2('13.5 Sécurité des paiements')
bullet("L\'intégration Stripe utilise le SDK officiel (Stripe Elements).")
bullet("Les données bancaires ne transitent jamais par le serveur backend — Stripe gère directement la tokenisation.")
bullet("Les webhooks Stripe sont vérifiés par signature HMAC.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  14. ENVIRONNEMENT DE TRAVAIL ET DÉPLOIEMENT
# ══════════════════════════════════════════════════════════════════════════════
h1('14. Environnement de Travail et Déploiement')

h2('14.1 Environnement de développement')
tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
table_header(tbl, ['Outil', 'Version / Édition', 'Usage'])
env_rows = [
    ('IntelliJ IDEA', 'Community 2024.x', 'Développement backend Spring Boot'),
    ('Visual Studio Code', '1.9x', 'Développement frontend React'),
    ('Git + GitHub', 'Desktop + web', 'Versioning, CI/CD, suivi de projet'),
    ('Postman', '10.x', 'Test des endpoints API REST'),
    ('pgAdmin 4', '7.x', 'Administration base de données PostgreSQL'),
    ('Node.js', '20 LTS', 'Exécution de Vite et du serveur de dev React'),
    ('Maven', '3.9', 'Build et gestion des dépendances Java'),
    ('PlantUML', 'Web / plugin IDE', 'Génération des diagrammes UML'),
]
for row in env_rows:
    add_row(tbl, row)

doc.add_paragraph()
h2('14.2 Infrastructure de production')
body("L\'application est déployée entièrement dans le cloud, sans serveur physique :")

tbl2 = doc.add_table(rows=1, cols=4)
tbl2.style = 'Table Grid'
table_header(tbl2, ['Composant', 'Plateforme', 'Localisation', 'Caractéristiques'])
infra_rows = [
    ('Base de données PostgreSQL 16', 'Neon.tech', 'EU-Central (Frankfurt)', 'Serverless, scaling automatique, backups quotidiens, RGPD compliant'),
    ('Backend Spring Boot', 'Render', 'Oregon (US West)', 'Conteneur Linux, HTTPS natif, logs temps réel, auto-deploy GitHub'),
    ('Frontend React/Vite', 'Vercel', 'CDN mondial (edge)', 'CDN global, HTTPS natif, preview deployments, auto-deploy GitHub'),
    ('Chatbot IA', 'Groq Cloud', 'US', 'LPU inference, modèle LLaMA 3.1-8b-instant, latence < 1s'),
    ('TTS vocal', 'ElevenLabs', 'Cloud', 'eleven_turbo_v2_5, voix Adam (fr), API REST'),
    ('Paiements', 'Stripe', 'Cloud', 'Mode test/production, 3D Secure, webhooks'),
]
for row in infra_rows:
    add_row(tbl2, row)

doc.add_paragraph()
h2('14.3 Pipeline CI/CD')
body("Le déploiement continu est automatisé via les intégrations GitHub des plateformes cloud :")
numbered("Un développeur pousse du code sur la branche main de GitHub.")
numbered("GitHub notifie Render et Vercel via webhooks.")
numbered("Render déclenche : mvn clean package -DskipTests, puis démarrage du JAR.")
numbered("Vercel déclenche : npm run build (Vite), puis déploiement du dist/ sur le CDN.")
numbered("L\'application est en production en moins de 3 minutes.")

h2('14.4 Variables d\'environnement')
body("Les secrets et configurations sont injectés au démarrage via des variables d\'environnement (jamais dans le code) :")
tbl3 = doc.add_table(rows=1, cols=3)
tbl3.style = 'Table Grid'
table_header(tbl3, ['Variable', 'Couche', 'Description'])
env_vars = [
    ('SPRING_DATASOURCE_URL', 'Backend', 'URL JDBC de la base Neon.tech'),
    ('SPRING_DATASOURCE_USERNAME', 'Backend', 'Utilisateur PostgreSQL'),
    ('SPRING_DATASOURCE_PASSWORD', 'Backend', 'Mot de passe PostgreSQL'),
    ('STRIPE_SECRET_KEY', 'Backend', 'Clé secrète Stripe (sk_live_...)'),
    ('VITE_API_BASE_URL', 'Frontend', 'URL de l\'API backend'),
    ('VITE_GROQ_API_KEY', 'Frontend', 'Clé API Groq pour le chatbot'),
    ('VITE_ELEVENLABS_API_KEY', 'Frontend', 'Clé API ElevenLabs pour TTS'),
    ('VITE_STRIPE_PUBLISHABLE_KEY', 'Frontend', 'Clé publique Stripe (pk_live_...)'),
]
for row in env_vars:
    add_row(tbl3, row)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  15. ANALYSE DU CHATBOT VOCAL
# ══════════════════════════════════════════════════════════════════════════════
h1('15. Analyse Technique du Chatbot Vocal')
body(
    "Le chatbot vocal est une fonctionnalité différenciante du projet, combinant trois technologies "
    "spécialisées pour offrir une expérience conversationnelle naturelle en français."
)

h2('15.1 Architecture du chatbot vocal')
tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
table_header(tbl, ['Composant', 'Technologie', 'Rôle'])
chatbot_rows = [
    ('Reconnaissance vocale (STT)', 'Web Speech API (navigateur Chrome)', 'Convertir la parole en texte en temps réel, langue fr-FR'),
    ('Intelligence artificielle (NLU)', 'Groq API — modèle LLaMA 3.1-8b-instant', 'Générer une réponse contextuelle au texte transcrit'),
    ('Synthèse vocale (TTS)', 'ElevenLabs API — modèle eleven_turbo_v2_5', 'Convertir le texte de réponse en audio naturel français'),
    ('Fallback TTS', 'Web Speech Synthesis API (navigateur)', 'Lire la réponse si ElevenLabs est indisponible'),
    ('Gestion d\'état', 'React useState/useRef', 'Gérer les transitions LISTENING → PROCESSING → SPEAKING'),
]
for row in chatbot_rows:
    add_row(tbl, row)

doc.add_paragraph()
h2('15.2 Machine à états du chatbot')
body("Le chatbot suit un cycle d\'état strict pour éviter les conflits audio/micro :")
bullet("IDLE : chatbot inactif, microphone désactivé.")
bullet("LISTENING : microphone actif, reconnaissance vocale en cours (Web Speech API).")
bullet("PROCESSING : texte reçu, requête en cours vers Groq API.")
bullet("SPEAKING : réponse reçue, lecture audio en cours (ElevenLabs / speechSynthesis).")
body("Transition automatique SPEAKING → LISTENING à la fin de l\'audio (événement onended).")
body("Protection anti-écho : le microphone est désactivé pendant la lecture TTS pour éviter que le chatbot s\'entende lui-même.")

h2('15.3 Optimisations de latence')
bullet("ElevenLabs turbo v2.5 : latence réduite de 50% par rapport aux modèles standards.")
bullet("Découpage de la réponse en phrases courtes pour un streaming perçu plus naturel.")
bullet("Timer de sécurité (15s) : coupe l\'écoute si aucune parole détectée.")
bullet("Timer global (30s) : réinitialise le chatbot en cas de blocage d\'état.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  16. PLANIFICATION
# ══════════════════════════════════════════════════════════════════════════════
h1('16. Planification du Développement')

h2('16.1 Découpage par phases')
tbl = doc.add_table(rows=1, cols=4)
tbl.style = 'Table Grid'
table_header(tbl, ['Phase', 'Description', 'Durée estimée', 'Livrables'])
phases = [
    ('Phase 1 — Fondations',
     'Mise en place infrastructure, modèle de données, authentification, CRUD de base',
     '3 semaines',
     'BD PostgreSQL, API auth, gestion utilisateurs'),
    ('Phase 2 — Fonctionnalités core',
     'Commandes, stock, caisse, tables, réservations',
     '4 semaines',
     'API commandes/stock/paiements, interface serveur/caissier/cuisinier'),
    ('Phase 3 — Temps réel & IA',
     'WebSocket, alertes stock, chatbot textuel, chatbot vocal',
     '3 semaines',
     'WebSocket STOMP, intégration Groq, ElevenLabs TTS'),
    ('Phase 4 — Client & Fournisseur',
     'Espace client (menu, commandes en ligne, Stripe, avis), espace fournisseur',
     '2 semaines',
     'Interface client, paiement Stripe, réservations en ligne'),
    ('Phase 5 — Analytics & Finition',
     'Tableau de bord, audit log, export PDF, tests, déploiement cloud',
     '2 semaines',
     'Dashboard, jsPDF, déploiement Render/Vercel'),
]
for row in phases:
    add_row(tbl, row)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  17. CONCLUSION ET PERSPECTIVES
# ══════════════════════════════════════════════════════════════════════════════
h1('17. Conclusion et Perspectives')

h2('17.1 Synthèse')
body(
    "Ce cahier d\'analyse définit une architecture cohérente, moderne et déployable pour l\'application "
    "Gestion Snack. L\'analyse couvre l\'ensemble des 6 rôles métier, les 26 user stories, les diagrammes "
    "UML complets (cas d\'utilisation, classes, 5 séquences), le modèle de données avec ses 14 tables, "
    "le dictionnaire de données exhaustif, les choix technologiques justifiés et l\'architecture de déploiement cloud."
)
body(
    "La valeur ajoutée du projet réside dans son chatbot vocal intelligent (Web Speech API + Groq/LLaMA 3.1 + "
    "ElevenLabs TTS), sa gestion automatique des stocks par triggers PostgreSQL, ses notifications temps réel "
    "par WebSocket et son déploiement cloud entièrement automatisé."
)

h2('17.2 Perspectives d\'évolution')
body("Les évolutions envisagées pour les versions futures du système incluent :")
bullet("Authentification multi-facteurs (2FA) via email ou application TOTP.")
bullet("Application mobile native (React Native) pour les serveurs et clients.")
bullet("Système de fidélité avec points et récompenses pour les clients réguliers.")
bullet("Intégration d\'un système de réservation avec confirmation par SMS/email.")
bullet("Tableau de bord prédictif avec machine learning pour anticiper les ruptures de stock.")
bullet("Support multi-établissements pour les chaînes de snacks.")
bullet("Mode hors-ligne (PWA) pour fonctionner sans connexion internet.")

h2('17.3 Conformité aux exigences TFE')
tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
table_header(tbl, ['Exigence TFE', 'Couverture', 'Section'])
tfe_rows = [
    ('Cahier des charges', 'Rédigé (document séparé)', 'Document Cahier des Charges V3'),
    ('Étude de l\'existant', 'Analyse comparative de 4 solutions', 'Section 2'),
    ('User stories complètes', '26 user stories sur 6 rôles avec critères d\'acceptation', 'Section 4'),
    ('Choix technologiques justifiés', '21 technologies justifiées individuellement', 'Section 7'),
    ('Analyse technique approfondie', 'UML complet, MCD, dictionnaire, sécurité, séquences', 'Sections 8 à 15'),
    ('GitHub + intégration continue', 'CI/CD automatique Render + Vercel sur push main', 'Section 14.3'),
    ('Environnement de travail', 'Détaillé avec outils et versions', 'Section 14.1'),
    ('Déploiement sur serveur Linux', 'Render = conteneur Linux (répond à l\'exigence)', 'Section 14.2'),
    ('Démonstration des compétences', 'Architecture N-tier, IA, temps réel, sécurité', 'Ensemble du document'),
]
for row in tfe_rows:
    add_row(tbl, row)

doc.add_paragraph('\n\n')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('— Fin du Cahier d\'Analyse —')
r.font.size = Pt(11)
r.font.italic = True
r.font.color.rgb = TEAL

# ══════════════════════════════════════════════════════════════════════════════
#  SAUVEGARDE
# ══════════════════════════════════════════════════════════════════════════════
output_path = r"d:\Projet TFE\gestion-snack\fichiers\Cahier_danalyse_Gestion_Snack_V3_Final.docx"
doc.save(output_path)
print(f"[OK] Document généré : {output_path}")
